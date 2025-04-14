import os

import gemmi
from docx import Document
from docx.oxml.xmlchemy import BaseOxmlElement
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from lxml import etree
from lxml.etree import XSLTAccessControl

from finalcif.app_path import application_path
from finalcif.cif.cif_file_io import CifContainer
from finalcif.cif.text import retranslate_delimiter, string_to_utf8
from finalcif.report import references
from finalcif.tools.misc import (protected_space, angstrom, remove_line_endings)


def math_to_word(eq: str) -> BaseOxmlElement:
    """Transform a sympy equation to be printed in word document."""
    tree = etree.fromstring(eq)
    xslt = etree.parse(os.path.join(application_path, 'template/mathml2omml.xsl'))
    acess_control = XSLTAccessControl(read_network=False, write_network=False)
    transform = etree.XSLT(xslt, access_control=acess_control)
    new_dom = transform(tree)
    return new_dom.getroot()


def clean_string(string: str) -> str:
    """
    Removes control characters from a string.
    """
    repl = string.replace('\t', ' ') \
        .replace('\f', ' ') \
        .replace('\0', ' ') \
        .strip(' ') \
        .strip('.')
    return remove_line_endings(repl)


def gstr(string: str) -> str:
    """
    Turn a string into a gemmi string and remove control characters.
    """
    return clean_string(gemmi.cif.as_string(string))


def _get_cooling_device(cif: CifContainer) -> str:
    olx = gstr(cif['_olex2_diffrn_ambient_temperature_device'])
    iucr = gstr(cif['_diffrn_measurement_ambient_temperature_device_make'])
    if olx and iucr:
        return iucr
    if olx:
        return olx
    elif iucr:
        return iucr
    else:
        return ''


class FormatMixin:

    def bold(self, run: Run):
        r = run.bold = True
        return r


class SpaceChar:
    def __init__(self, paragraph: Paragraph):
        self.p = paragraph

    def regular(self) -> None:
        self.p.add_run(' ')

    def protected(self):
        self.p.add_run(protected_space)


class RefinementDetails:
    def __init__(self, cif: CifContainer, document: Document):
        ph = document.add_paragraph(style='Heading 2')
        ph.add_run(text=fr"Refinement details for {cif.block.name}")
        p = document.add_paragraph()
        try:
            p.style = document.styles['fliesstext']
        except KeyError:
            print('DBG> Text style not found')
        text = ' '.join(cif['_refine_special_details'].splitlines(keepends=False))
        if cif['_olex2_refine_details']:
            text += ' '.join(cif['_olex2_refine_details'].splitlines(keepends=False))
        # Replacing semicolon, because it can damage the CIF:
        text = text.replace(';', '.')
        p.add_run(string_to_utf8(text).strip())


def get_inf_article(next_word: str) -> str:
    if not next_word:
        return 'a'
    voc = 'aeiou'
    return 'an' if next_word[0].lower() in voc else 'a'


def get_distance_unit(picometers: bool) -> str:
    if picometers:
        return 'pm'
    else:
        return angstrom


def get_volume_unit(picometers: bool) -> str:
    if picometers:
        return 'nm'
    else:
        return angstrom


def format_float_with_decimal_places(number: float, places=2) -> str:
    try:
        fnum = float(number)
        return f'{fnum:.{places}f}'
    except ValueError:
        return f'{number}'


def utf8(text: str) -> str:
    """A Jinja2 filter for CIF to utf8"""
    return string_to_utf8(text)


def format_radiation(radiation_type: str) -> list:
    radtype = list(radiation_type.partition("K"))
    if len(radtype) > 2:
        radtype[2] = retranslate_delimiter(radtype[2])
        return radtype
    else:
        return radtype


def make_report_text(cif, document: Document) -> references.ReferenceList:
    paragr = document.add_paragraph()
    try:
        paragr.style = document.styles['fliesstext']
    except KeyError:
        print('DBG> Text style not found')
    ref = references.ReferenceList(paragr)
    return ref
