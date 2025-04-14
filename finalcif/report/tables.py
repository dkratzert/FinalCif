#
# This program was innitially created by Nils Trapp and extended by
# Daniel Kratzert
#
import itertools as it
from collections.abc import Sequence
from math import sin, radians
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Length, Pt
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from finalcif.app_path import application_path
from finalcif.cif.cif_file_io import CifContainer
from finalcif.report.mtools import format_space_group
from finalcif.report.references import ReferenceList
from finalcif.report.report_text import format_radiation, RefinementDetails, make_report_text
from finalcif.report.templated_report import BondsAndAngles, TorsionAngles, HydrogenBonds
from finalcif.tools.misc import (protected_space, angstrom, bequal, sigma_sm, halbgeviert, degree_sign,
                                 ellipsis_mid, less_or_equal, timessym, lambdasym, this_or_quest,
                                 isnumeric, minus_sign, theta_symbol, grouper, open_file, pi_symbol,
                                 ellipsis_char, medium_math_space, alpha_symol, beta_symol,
                                 gamma_symol)
from finalcif.tools.options import Options
from finalcif.tools.settings import FinalCifSettings


def open_cif_file(cif_fileobj) -> None | CifContainer:
    try:
        cif = CifContainer(cif_fileobj)
    except Exception as e:
        print('Unable to open file', cif_fileobj.name)
        print(e)
        return None
    return cif


def make_multi_tables(cif: CifContainer, output_filename: str = 'multitable.docx') -> str:
    """
    The main loop for doing the report pages with tables.
    """
    group_of_files = list(grouper([x for x in cif.doc if x.name != 'global'], 3))
    table_index = len(group_of_files) - 1
    document = create_document()
    document.add_heading('Structure Tables', 1)
    for file_group in enumerate(group_of_files):
        # page_number = file_group[0]
        cif_triple = file_group[1]
        main_table = document.add_table(rows=1, cols=4)
        populate_description_columns(main_table, cif, row_shift=1)
        for table_column in range(1, 4):  # the three columns
            if cif_triple[table_column - 1]:
                cif_block = file_group[1][table_column - 1]
                cif.block = cif_block
                populate_main_table_values(main_table, cif, column=table_column, row_shift=1)
        if file_group[0] < table_index:
            document.add_page_break()

    print('\nScript finished - output file: multitable.docx')
    document.save(output_filename)
    return output_filename


def make_report_from(options: Options, cif: CifContainer, output_filename: str | None = None, picfile: Path | None = None) -> str:
    """
    Creates a tabular cif report.
    """
    document = create_document()
    references: ReferenceList | None = None

    if not options.report_text:
        document.add_heading(f'Structure Tables for {cif.block.name}', 1)
    else:
        document.add_heading('Structure Tables', 1)
        make_columns_section(document, columns='2')

    if options.report_text:
        if picfile and picfile.exists():
            add_picture(document, options, picfile)
        references = make_report_text(cif, document)

    # -- The residuals table:
    table_num = 1
    if options.report_text:
        # I have to do the header and styling here, otherwise I get another paragraph
        # with a line break in front of the heading.
        p = document.add_paragraph(style='Heading 2')
        p.add_run().add_break(WD_BREAK.COLUMN)
        tab0_head = fr"Table {table_num}. Crystal data and structure refinement for {cif.block.name}"
        p.add_run(text=tab0_head)
    table_num = add_residuals_table(document, cif, table_num)
    p = document.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    if options.report_text:
        make_columns_section(document, columns='1')

    if ((cif['_refine_special_details'].strip() != '' or cif['_olex2_refine_details'].strip() != '') and not
    (cif['_refine_special_details'] == '?' or cif['_olex2_refine_details'] == '?') and options.report_text):
        RefinementDetails(cif, document)

    table_num = add_coords_table(document, cif, table_num)
    if options.report_adp and len(tuple(cif.displacement_parameters())) > 0:
        table_num = add_adp_table(document, cif, table_num)

    if cif.symmops:
        if len(list(cif.bonds())) + len(list(cif.angles())) > 0:
            table_num += 1
            document.add_heading(fr"Table {table_num}. Bond lengths and angles for {cif.block.name}", 2)
            make_columns_section(document, columns='2')
            ba = BondsAndAngles(cif, without_h=options.without_h)
            table_num = add_bonds_and_angles_table(document, table_num, ba)
        if len(list(cif.torsion_angles())) > 0:
            make_columns_section(document, columns='1')
            table_num += 1
            document.add_heading(fr"Table {table_num}. Torsion angles for {cif.block.name}", 2)
            make_columns_section(document, columns='2')
            tors = TorsionAngles(cif, without_h=options.without_h)
            table_num = add_torsion_angles(document, table_num, tors)
        make_columns_section(document, columns='1')
        if len(list(cif.hydrogen_bonds())) > 0:
            table_num += 1
            h = HydrogenBonds(cif)
            document.add_heading(fr"Table {table_num}. Hydrogen bonds for {cif.block.name}", 2)
            table_num = add_hydrogen_bonds(document, table_num, h)
        document.add_paragraph('')
    else:
        make_columns_section(document, columns='1')
        document.add_paragraph('No further tables, because symmetry operators '
                               '(_space_group_symop_operation_xyz) are missing.')
    if options.report_text:
        # -- Bibliography:
        document.add_heading('Bibliography', 2)
        references.make_literature_list(document)

    document.save(output_filename)
    print(f'\nTables finished - output file: {output_filename}')
    return cif.block.name


def add_picture(document: Document, options: Options, picfile: Path) -> None:
    pic = document.add_paragraph()
    try:
        width = float(options.picture_width)
    except ValueError:
        width = 7.0
    pic.add_run().add_picture(str(picfile), width=Cm(width))


def create_document() -> Document:
    """
    Creates the report docx document.

    :return: The document instance.
    """
    try:
        document = Document(Path(application_path).joinpath('template/template1.docx').resolve().__str__())
    except FileNotFoundError as e:
        print(e)
        document = Document()
    # Deleting first (empty) paragraph, otherwise first line would be an empty one:
    delete_first_paragraph(document)
    return document


def delete_paragraph(paragraph: Paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def delete_first_paragraph(document: Document):
    try:
        p = document.paragraphs[0]
        delete_paragraph(p)
    except IndexError:
        # no paragraph there
        pass


def make_columns_section(document, columns: str = '1'):
    """
    Makes a new section (new page) which has a certain number of columns.
    available sections:
    CONTINUOUS, NEW_COLUMN, NEW_PAGE, EVEN_PAGE, ODD_PAGE
    """
    # noinspection PyUnresolvedReferences
    from docx.enum.section import WD_SECTION
    section = document.add_section(WD_SECTION.CONTINUOUS)
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), f'{columns}')


def set_cell_border(cell: _Cell, **kwargs) -> None:
    """
    Set cell`s border of a Table cell instance.

    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = f'w:{edge}'
            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn(f'w:{key}'), str(edge_data[key]))


def add_residuals_table(document: Document, cif: CifContainer, table_num: int) -> int:
    main_table = document.add_table(rows=0, cols=2)
    # setup table format:
    set_column_width(main_table.columns[0], Cm(4.05))
    set_column_width(main_table.columns[1], Cm(4.05))
    # Add descriptions to the first column of the main table:
    populate_description_columns(main_table, cif)
    # The main residuals table:
    populate_main_table_values(main_table, cif)
    return table_num


def populate_description_columns(main_table: Table, cif: CifContainer, row_shift=0) -> None:
    """
    This Method adds the descriptions to the fist property table column.
    row_shift: Moves alle cells down by this factor to allow an empty first row for the data block name.
    """
    # paragraph(main_table).add_run('')
    paragraph(main_table).add_run('CCDC number')
    paragraph(main_table).add_run('Empirical formula')
    paragraph(main_table).add_run('Formula weight')
    paragraph(main_table).add_run('Temperature [K]')
    paragraph(main_table).add_run('Crystal system')
    paragraph(main_table).add_run('Space group (number)')
    _a_b_c(main_table)
    paragraph(main_table).add_run(f'{alpha_symol} [{degree_sign}]')
    paragraph(main_table).add_run(f'{beta_symol} [{degree_sign}]')
    paragraph(main_table).add_run(f'{gamma_symol} [{degree_sign}]')
    _volume(main_table)
    _rho(main_table)
    _mu(main_table)
    _crystal_size(main_table)
    paragraph(main_table).add_run('Crystal colour')
    paragraph(main_table).add_run('Crystal shape')
    paragraph(main_table).add_run('Radiation')
    paragraph(main_table).add_run(f'2{theta_symbol} range [{degree_sign}]')
    paragraph(main_table).add_run('Index ranges')
    paragraph(main_table).add_run('Reflections collected')
    paragraph(main_table).add_run('Independent reflections')
    _completeness(cif, main_table)
    paragraph(main_table).add_run('Data / Restraints / Parameters')
    _absotion_correction(main_table)
    _goof(main_table)
    # paragraph(main_table).add_run('Weigthing Scheme')
    _r_values_cutoff(main_table)
    _r_values(main_table)
    _peaks(main_table)
    _flack_x(main_table, cif, row_shift)
    _extinction(main_table, cif, row_shift)


def populate_main_table_values(main_table: Table, cif: CifContainer, column=1, row_shift=0):
    """
    Fills the main table with residuals. Column, by column.
    """
    row = 0
    if row_shift > 0:
        main_table.cell(row, column).paragraphs[0].add_run(cif.block.name).bold = True
        row += 1
    main_table.cell(row, column).paragraphs[0].add_run(cif['_database_code_depnum_ccdc_archive'])
    row += 1
    formula_paragraph = main_table.cell(row, column).paragraphs[0]
    sum_formula = cif['_chemical_formula_sum'].replace(" ", "")
    add_sum_formula(formula_paragraph, sum_formula)
    row += 1
    add_table_value(cif, main_table, '_chemical_formula_weight', row, column)
    row += 1
    add_table_value(cif, main_table, '_diffrn_ambient_temperature', row, column)
    row += 1
    add_table_value(cif, main_table, '_space_group_crystal_system', row, column)
    row += 1
    spgr_paragraph = main_table.cell(row, column).paragraphs[0]
    space_group = cif.space_group
    try:
        it_number = str(cif.spgr_number)
    except AttributeError:
        it_number = ''
    format_space_group(spgr_paragraph, space_group, it_number)
    row += 1
    add_table_value(cif, main_table, '_cell_length_a', row, column)
    row += 1
    add_table_value(cif, main_table, '_cell_length_b', row, column)
    row += 1
    add_table_value(cif, main_table, '_cell_length_c', row, column)
    row += 1
    add_table_value(cif, main_table, '_cell_angle_alpha', row, column)
    row += 1
    add_table_value(cif, main_table, '_cell_angle_beta', row, column)
    row += 1
    add_table_value(cif, main_table, '_cell_angle_gamma', row, column)
    row += 1
    add_table_value(cif, main_table, '_cell_volume', row, column)
    row += 1
    add_table_value(cif, main_table, '_cell_formula_units_Z', row, column)
    row += 1
    add_table_value(cif, main_table, '_exptl_crystal_density_diffrn', row, column)
    row += 1
    add_table_value(cif, main_table, '_exptl_absorpt_coefficient_mu', row, column)
    row += 1
    add_table_value(cif, main_table, '_exptl_crystal_F_000', row, column)
    try:
        completeness = f"{round(float(cif['_diffrn_measured_fraction_theta_full']) * 100, 1):.1f} %"
    except ValueError:
        completeness = '?'
    row += 1
    crystal_size_min = cif['_exptl_crystal_size_min']
    crystal_size_mid = cif['_exptl_crystal_size_mid']
    crystal_size_max = cif['_exptl_crystal_size_max']
    main_table.cell(row, column).text = (f'{this_or_quest(crystal_size_max)}{timessym}'
                                         f'{this_or_quest(crystal_size_mid)}{timessym}'
                                         f'{this_or_quest(crystal_size_min)}')
    row += 1
    add_table_value(cif, main_table, '_exptl_crystal_colour', row, column)
    row += 1
    add_table_value(cif, main_table, '_exptl_crystal_description', row, column)
    row += 1
    radiation_wavelength = cif['_diffrn_radiation_wavelength']
    wavelength = str(f' ({lambdasym} = {this_or_quest(radiation_wavelength)}'
                     f'{protected_space}{angstrom})').replace(' ', '')
    # radtype: ('Mo', 'K', '\\a')
    radiation_type = cif['_diffrn_radiation_type']
    radtype = format_radiation(radiation_type)
    radrun = main_table.cell(row, column).paragraphs[0]
    # radiation type e.g. Mo:
    radrun.add_run(radtype[0])
    # K line:
    radrunita = radrun.add_run(radtype[1])
    radrunita.font.italic = True
    alpha = radrun.add_run(radtype[2])
    alpha.font.italic = True
    alpha.font.subscript = True
    # wavelength lambda:
    radrun.add_run(' ' + wavelength)
    row += 1
    theta_min = cif['_diffrn_reflns_theta_min']
    theta_max = cif['_diffrn_reflns_theta_max']
    try:
        d_max = f' ({float(radiation_wavelength) / (2 * sin(radians(float(theta_max)))):.2f}{protected_space}{angstrom})'
        # 2theta range:
        main_table.cell(row,
                        column).text = f"{2 * float(theta_min):.2f} to {2 * float(theta_max):.2f}{d_max}"
    except ValueError:
        main_table.cell(row, column).text = '? to ?'
    row += 1
    add_hkl_indices(main_table, cif, row, column)
    row += 1
    add_table_value(cif, main_table, '_diffrn_reflns_number', row, column)
    row += 1
    rint_p = main_table.cell(row, column).paragraphs[0]
    add_r_int_value(cif, rint_p)
    row += 1
    main_table.cell(row, column).paragraphs[0].add_run(completeness)
    row += 1
    ls_number_reflns = cif['_refine_ls_number_reflns']
    ls_number_restraints = cif['_refine_ls_number_restraints']
    ls_number_parameters = cif['_refine_ls_number_parameters']
    main_table.cell(row, column).text = (f'{this_or_quest(ls_number_reflns)}/'
                                         f'{this_or_quest(ls_number_restraints)}/'
                                         f'{this_or_quest(ls_number_parameters)}')
    row += 1
    main_table.cell(row, column).text = (f"{cif['_exptl_absorpt_correction_T_min'] or '?'}/"
                                         f"{cif['_exptl_absorpt_correction_T_max'] or '?'}\n"
                                         f"({cif['_exptl_absorpt_correction_type'] or '?'})")
    row += 1
    add_table_value(cif, main_table, '_refine_ls_goodness_of_fit_ref', row, column)
    row += 1
    # add_table_value(cif, main_table, '_refine_ls_weighting_details', row, column)
    # row += 1
    add_r1sig_and_wr2full(main_table, cif, row, column)
    row += 2
    try:
        diff_density_min = f"{round(float(cif['_refine_diff_density_min']), 2):.2f}".replace('-', minus_sign)
    except ValueError:
        diff_density_min = '?'
    try:
        diff_density_max = f"{round(float(cif['_refine_diff_density_max']), 2):.2f}".replace('-', minus_sign)
    except ValueError:
        diff_density_max = '?'
    main_table.cell(row, column).text = f'{diff_density_max}/{diff_density_min}'
    row += 1
    if not cif.is_centrosymm:
        main_table.cell(row, column).text = cif['_refine_ls_abs_structure_Flack'].replace('-', minus_sign) or '?'
        row += 1
    elif row_shift > 0:
        main_table.cell(row, column).text = '---'
        row += 1
    exti = cif['_refine_ls_extinction_coef']
    if exti:
        main_table.cell(row, column).text = exti
    elif row_shift > 0:
        main_table.cell(row, column).text = '---'
    row += 1


def _absotion_correction(main_table: Table) -> None:
    p = paragraph(main_table)
    p.add_run('Absorption correction\nT')
    p.add_run('min').font.subscript = True
    p.add_run('/T')
    p.add_run('max').font.subscript = True
    p.add_run(' (method)')


def _extinction(main_table: Table, cif: CifContainer, row_shift: int) -> None:
    exti = cif['_refine_ls_extinction_coef']
    if exti or row_shift > 0:
        # always the last cell
        paragraph(main_table).add_run('Extinction coefficient')


def _flack_x(main_table: Table, cif: CifContainer, row_shift: int) -> None:
    if not cif.is_centrosymm or row_shift > 0:
        p = paragraph(main_table)
        p.add_run('Flack X parameter')


def _peaks(main_table: Table) -> None:
    p = paragraph(main_table)
    p.add_run(f'Largest peak/hole [e{angstrom}')
    p.add_run(minus_sign + '3').font.superscript = True
    p.add_run(']')


def _r_values(main_table: Table) -> None:
    p = paragraph(main_table)
    p.add_run('Final ')
    p.add_run('R').font.italic = True
    p.add_run(' indexes \n[all data]')


def _r_values_cutoff(main_table: Table) -> None:
    p = paragraph(main_table)
    p.add_run('Final ')
    p.add_run('R').font.italic = True
    p.add_run(' indexes \n[')
    p.add_run('I').font.italic = True
    p.add_run(f'{bequal}2{sigma_sm}(')
    p.add_run('I').font.italic = True
    p.add_run(')]')


def _goof(main_table: Table) -> None:
    p = paragraph(main_table)
    p.add_run('Goodness-of-fit on ')
    p.add_run('F').font.italic = True
    p.add_run('2').font.superscript = True


def _completeness(cif, main_table: Table) -> None:
    p = paragraph(main_table)
    theta_full = cif['_diffrn_reflns_theta_full']
    if theta_full:
        p.add_run(f'Completeness to \n{theta_symbol} = {theta_full}{degree_sign}')
    else:
        p.add_run('Completeness')


def _a_b_c(main_table: Table) -> None:
    p = paragraph(main_table)
    p.add_run('a').font.italic = True
    p.add_run(f' [{angstrom}]')
    p = paragraph(main_table)
    p.add_run('b').font.italic = True
    p.add_run(f' [{angstrom}]')
    p = paragraph(main_table)
    p.add_run('c').font.italic = True
    p.add_run(f' [{angstrom}]')


def _volume(main_table: Table) -> None:
    p = paragraph(main_table)
    p.add_run(f'Volume [{angstrom}')
    p.add_run('3').font.superscript = True
    p.add_run(']')
    paragraph(main_table).add_run('Z').font.italic = True


def _rho(main_table: Table) -> None:
    p = paragraph(main_table)
    p.add_run('\u03C1').font.italic = True  # rho
    p.add_run('calc').font.subscript = True
    p.add_run(' [gcm')
    p.add_run(minus_sign + '3').font.superscript = True
    p.add_run(']')


def _mu(main_table: Table) -> None:
    p = paragraph(main_table)
    p.add_run('\u03BC').font.italic = True  # mu
    p.add_run(' [mm')
    p.add_run(minus_sign + '1').font.superscript = True
    p.add_run(']')


def _crystal_size(main_table: Table) -> None:
    p = paragraph(main_table)
    p.add_run('F').font.italic = True
    p.add_run('(000)')
    p = paragraph(main_table)
    p.add_run('Crystal size [mm')
    p.add_run('3').font.superscript = True
    p.add_run(']')


def add_hkl_indices(main_table: Table, cif: CifContainer, row: int, column: int) -> None:
    if all([cif['_diffrn_reflns_limit_h_min'], cif['_diffrn_reflns_limit_h_max'],
            cif['_diffrn_reflns_limit_k_min'], cif['_diffrn_reflns_limit_k_max'],
            cif['_diffrn_reflns_limit_l_min'], cif['_diffrn_reflns_limit_l_max']
            ]):
        limit_h_min = cif['_diffrn_reflns_limit_h_min']
        limit_h_max = cif['_diffrn_reflns_limit_h_max']
        limit_k_min = cif['_diffrn_reflns_limit_k_min']
        limit_k_max = cif['_diffrn_reflns_limit_k_max']
        limit_l_min = cif['_diffrn_reflns_limit_l_min']
        limit_l_max = cif['_diffrn_reflns_limit_l_max']
    else:
        limit_h_min = '?'
        limit_h_max = '?'
        limit_k_min = '?'
        limit_k_max = '?'
        limit_l_min = '?'
        limit_l_max = '?'
    main_table.cell(row, column).text = (f'{minus_sign if limit_h_min != "0" else ""}'
                                         f'{limit_h_min.replace("-", "")} '
                                         f'{less_or_equal} h {less_or_equal} {limit_h_max}\n'
                                         f'{minus_sign if limit_k_min != "0" else ""}'
                                         f'{limit_k_min.replace("-", "")} '
                                         f'{less_or_equal} k {less_or_equal} {limit_k_max}\n'
                                         f'{minus_sign if limit_l_min != "0" else ""}'
                                         f'{limit_l_min.replace("-", "")} '
                                         f'{less_or_equal} l {less_or_equal} {limit_l_max}')


def add_r1sig_and_wr2full(main_table: Table, cif: CifContainer, row: int, column: int) -> None:
    r2sig_p = main_table.cell(row, column).paragraphs[0]
    rfull_p = main_table.cell(row + 1, column).paragraphs[0]
    ls_R_factor_gt = cif['_refine_ls_R_factor_gt']
    ls_wR_factor_gt = cif['_refine_ls_wR_factor_gt']
    ls_R_factor_all = cif['_refine_ls_R_factor_all']
    ls_wR_factor_ref = cif['_refine_ls_wR_factor_ref']

    r2sig_p.add_run('R').font.italic = True
    r2sig_p.add_run('1').font.subscript = True
    r2sig_p.add_run(' = ' + this_or_quest(ls_R_factor_gt))
    r2sig_p.add_run('\nw')
    r2sig_p.add_run('R').font.italic = True
    r2sig_p.add_run('2').font.subscript = True
    r2sig_p.add_run(' = ' + this_or_quest(ls_wR_factor_gt))

    rfull_p.add_run('R').font.italic = True
    rfull_p.add_run('1').font.subscript = True
    rfull_p.add_run(' = ' + this_or_quest(ls_R_factor_all))
    rfull_p.add_run('\nw')
    rfull_p.add_run('R').font.italic = True
    rfull_p.add_run('2').font.subscript = True
    rfull_p.add_run(' = ' + ls_wR_factor_ref)


def add_r_int_value(cif: CifContainer, rint_p: Paragraph):
    reflns_number_total = cif['_reflns_number_total']
    reflns_av_R_equivalents = cif['_diffrn_reflns_av_R_equivalents']
    reflns_av_unetI = cif['_diffrn_reflns_av_unetI/netI']
    rint_p.add_run(this_or_quest(reflns_number_total) + '\n')
    rint_p.add_run('R').font.italic = True
    rint_p.add_run('int').font.subscript = True
    rint_p.add_run(' = ' + this_or_quest(reflns_av_R_equivalents) + '\n')
    rint_p.add_run('R').font.italic = True
    rint_p.add_run('sigma').font.subscript = True
    rint_p.add_run(' = ' + this_or_quest(reflns_av_unetI))


def add_table_value(cif: CifContainer, main_table: Table, keyword: str, row: int, column=1) -> None:
    cell = main_table.cell(row, column)
    if cif[keyword]:
        cell.text = cif[keyword]
    else:
        cell.text = '?'


def add_sum_formula(formula_paragraph: Paragraph, sum_formula: str) -> None:
    if sum_formula:
        sum_formula_group = [''.join(x[1]) for x in it.groupby(sum_formula, lambda x: x.isalpha())]
        for _, word in enumerate(sum_formula_group):
            formrunsub = formula_paragraph.add_run(word)
            if isnumeric(word):
                formrunsub.font.subscript = True
    else:
        formula_paragraph.add_run('no sum formula')


def add_decimal_tab(num_string: str) -> str:
    """
    Adds a tab character in front of the decimal point in order to get proper alignment of the tabstops.
    """
    return '\t.'.join(num_string.split('.'))


def add_coords_table(document: Document, cif: CifContainer, table_num: int):
    """
    Adds the table with the atom coordinates.
    :param document: The current word document.
    :param cif: the cif object from CifContainer.
    :return: None
    """
    atoms = list(cif.atoms())
    table_num += 1
    headline = f"Table {table_num}. Atomic coordinates and "
    h = document.add_heading(headline, 2)
    h.add_run('U').font.italic = True
    h.add_run('eq').font.subscript = True
    h.add_run(f'{protected_space}[{angstrom}')
    h.add_run('2').font.superscript = True
    h.add_run(f'] for {cif.block.name}')
    coords_table = document.add_table(rows=len(atoms) + 1, cols=5, style='Table Grid')
    # Atom	x	y	z	U(eq)
    head_row = coords_table.rows[0]
    head_row.cells[0].paragraphs[0].add_run('Atom').bold = True
    px = head_row.cells[1].paragraphs[0]
    ar = px.add_run('x')
    ar.bold = True
    ar.italic = True
    py = head_row.cells[2].paragraphs[0]
    ar = py.add_run('y')
    ar.bold = True
    ar.italic = True
    pz = head_row.cells[3].paragraphs[0]
    ar = pz.add_run('z')
    ar.bold = True
    ar.italic = True
    pu = head_row.cells[4].paragraphs[0]
    ar = pu.add_run('U')
    ar.bold = True
    ar.italic = True
    ar2 = pu.add_run('eq')
    ar2.bold = True
    ar2.font.subscript = True
    # having a list of column cells before is *much* faster!
    col0_cells = coords_table.columns[0].cells
    col1_cells = coords_table.columns[1].cells
    col2_cells = coords_table.columns[2].cells
    col3_cells = coords_table.columns[3].cells
    col4_cells = coords_table.columns[4].cells
    rowidx = 1
    for at in atoms:
        c0, c1, c2, c3, c4 = col0_cells[rowidx], col1_cells[rowidx], col2_cells[rowidx], \
            col3_cells[rowidx], col4_cells[rowidx]
        rowidx += 1
        c0.text = at[0]  # label
        c1.text = (str(at[2]))  # x
        c2.text = (str(at[3]))  # y
        c3.text = (str(at[4]))  # z
        c4.text = (str(at[7]))  # ueq
    p = document.add_paragraph()
    p.style = document.styles['tabunterschr']
    p.add_run('U').font.italic = True
    p.add_run('eq').font.subscript = True
    p.add_run(' is defined as 1/3 of the trace of the orthogonalized ')
    p.add_run('U').font.italic = True
    ij = p.add_run('ij')
    ij.font.subscript = True
    ij.font.italic = True
    p.add_run(' tensor.')
    set_column_width(coords_table.columns[0], Cm(2.3))
    set_column_width(coords_table.columns[1], Cm(2.8))
    set_column_width(coords_table.columns[2], Cm(2.8))
    set_column_width(coords_table.columns[3], Cm(2.8))
    set_column_width(coords_table.columns[4], Cm(2.8))
    document.add_paragraph()
    return table_num


def add_adp_table(document: Document, cif: CifContainer, table_num: int):
    """
    Anisotropic displacement parameters (Å2) for {{ cif.block.name }}.
    The anisotropic displacement factor exponent takes the form: -2π2[h2a*2U11 + ... + 2hka*b*U12]
    """
    table_num += 1
    headline = f"Table {table_num}. Anisotropic displacement parameters [{angstrom}"
    h = document.add_heading(headline, 2)
    h.add_run('2').font.superscript = True
    h.add_run(f'] for {cif.block.name}.\nThe anisotropic displacement factor exponent takes '
              f'the form: {minus_sign}2{pi_symbol}')
    h.add_run('2').font.superscript = True
    h.add_run('[')

    h.add_run(f'{medium_math_space}h').font.italic = True
    h.add_run('2').font.superscript = True
    h.add_run('(a*)').font.italic = True
    h.add_run('2').font.superscript = True
    h.add_run('U').font.italic = True
    h.add_run('11').font.subscript = True

    h.add_run(f'{medium_math_space}+{medium_math_space}k').font.italic = True
    h.add_run('2').font.superscript = True
    h.add_run('(b*)').font.italic = True
    h.add_run('2').font.superscript = True
    h.add_run('U').font.italic = True
    h.add_run('22').font.subscript = True

    h.add_run(f'{medium_math_space}+{medium_math_space}{ellipsis_char}'
              f'{medium_math_space}+{medium_math_space}')
    h.add_run('2hka*b*U').font.italic = True
    h.add_run('12').font.subscript = True
    h.add_run(f'{medium_math_space}]')

    adp_data = tuple(cif.displacement_parameters())
    adp_table = document.add_table(rows=len(adp_data) + 1, cols=7, style='Table Grid')
    head_row = adp_table.rows[0]
    head_row.cells[0].paragraphs[0].add_run('Atom').bold = True
    for n, u_val in enumerate(('11', '22', '33', '23', '13', '12'), 1):
        p_u = head_row.cells[n].paragraphs[0]
        u_char(p_u)
        u_number(p_u, u_val)
    col1_cells = adp_table.columns[0].cells
    col2_cells = adp_table.columns[1].cells
    col3_cells = adp_table.columns[2].cells
    col4_cells = adp_table.columns[3].cells
    col5_cells = adp_table.columns[4].cells
    col6_cells = adp_table.columns[5].cells
    col7_cells = adp_table.columns[6].cells
    for rowidx, row in enumerate(adp_data, 1):
        col1_cells[rowidx].text = row.label
        col2_cells[rowidx].text = row.U11.replace('-', minus_sign)
        col3_cells[rowidx].text = row.U22.replace('-', minus_sign)
        col4_cells[rowidx].text = row.U33.replace('-', minus_sign)
        col5_cells[rowidx].text = row.U23.replace('-', minus_sign)
        col6_cells[rowidx].text = row.U13.replace('-', minus_sign)
        col7_cells[rowidx].text = row.U12.replace('-', minus_sign)
    set_column_width(adp_table.columns[0], Cm(1.6))
    for num in range(1, 7):
        set_column_width(adp_table.columns[num], Cm(2.4))
    return table_num


def u_char(p_u):
    r = p_u.add_run('U')
    r.font.bold = True
    r.font.italic = True


def u_number(p_u, value):
    r = p_u.add_run(value)
    r.font.bold = True
    r.font.subscript = True


def add_bonds_and_angles_table(document: Document, table_num: int, data: BondsAndAngles = None) -> int:
    """
    Make table with bonds and angles.
    """
    # creating rows in advance is *much* faster!
    bond_angle_table = document.add_table(rows=len(data) + 3, cols=2, style='Table Grid')
    head_row = bond_angle_table.rows[0]
    head_row.cells[0].paragraphs[0].add_run(f'Atom{halbgeviert}Atom').bold = True
    p_length = head_row.cells[1].paragraphs[0]
    p_length.add_run(f'Length [{angstrom}]').bold = True
    # having a list of column cells before is *much* faster!
    col1_cells = bond_angle_table.columns[0].cells
    col2_cells = bond_angle_table.columns[1].cells
    rowidx = 1
    # filling rows:
    for b in data.bonds_as_string:
        c0, c1 = col1_cells[rowidx], col2_cells[rowidx]
        rowidx += 1
        c0.text = b.atoms
        c0.paragraphs[0].add_run(b.symm).font.superscript = True
        c1.text = str(b.dist)  # bond
    ############ the angles ####################
    head_row = bond_angle_table.rows[rowidx + 1]
    head_row.cells[0].paragraphs[0].add_run(f'Atom{halbgeviert}Atom{halbgeviert}Atom').bold = True
    head_row.cells[1].paragraphs[0].add_run(f'Angle [{degree_sign}]').bold = True
    set_cell_border(head_row.cells[0], bottom={"sz": 2, "color": "#000000", "val": "single"})
    set_cell_border(head_row.cells[1], bottom={"sz": 2, "color": "#000000", "val": "single"})
    rowidx += 2
    for a in data.angles_as_string:
        c0, c1 = col1_cells[rowidx], col2_cells[rowidx]
        rowidx += 1
        # atom1 symm1_str atom2 symm2_str
        c0.text = a.atom1
        cp0 = c0.paragraphs[0]
        cp0.add_run(a.symm1).font.superscript = True
        cp0.add_run(a.atom2)
        cp0.add_run(a.symm2).font.superscript = True
        c1.text = a.angle
    set_column_width(bond_angle_table.columns[0], Cm(3.7))
    set_column_width(bond_angle_table.columns[1], Cm(2.5))
    if data.without_h:
        add_hydrogen_omit_info(document)
    if data.symmetry_generated_atoms_used:
        add_last_symminfo_line(data.symminfo, document)
    return table_num


def add_torsion_angles(document: Document, table_num: int, data: TorsionAngles = None):
    """
    Table 6.  Torsion angles [°] for I-43d_final.
    """
    torsion_table = document.add_table(rows=len(data) + 1, cols=2, style='Table Grid')
    head_row = torsion_table.rows[0]
    head_row.cells[0].paragraphs[0].add_run(f'Atom{halbgeviert}Atom{halbgeviert}Atom{halbgeviert}Atom').bold = True
    head_row.cells[1].paragraphs[0].add_run(f'Torsion Angle [{degree_sign}]').bold = True
    col0_cells = torsion_table.columns[0].cells
    col1_cells = torsion_table.columns[1].cells
    rowidx = 1
    for tor in data.torsion_angles_as_string:
        c0, c1 = col0_cells[rowidx], col1_cells[rowidx]
        rowidx += 1
        cp0 = c0.paragraphs[0]
        cp0.add_run(tor.atom1)
        cp0.add_run(tor.symm1).font.superscript = True
        cp0.add_run(halbgeviert)
        cp0.add_run(tor.atom2)
        cp0.add_run(tor.symm2).font.superscript = True
        cp0.add_run(halbgeviert)
        cp0.add_run(tor.atom3)
        cp0.add_run(tor.symm3).font.superscript = True
        cp0.add_run(halbgeviert)
        cp0.add_run(tor.atom4)  # labels
        cp0.add_run(tor.symm4).font.superscript = True
        c1.paragraphs[0].add_run(tor.angle)  # torsion angle
    set_column_width(torsion_table.columns[0], Cm(4.2))
    set_column_width(torsion_table.columns[1], Cm(3.0))
    if data.without_h:
        add_hydrogen_omit_info(document)
    if data.symmetry_generated_atoms_used:
        add_last_symminfo_line(data.symminfo, document)
    return table_num


def add_hydrogen_bonds(document: Document, table_num: int, data: HydrogenBonds = None) -> int:
    """
    Table 7.  Hydrogen bonds for I-43d_final  [Å and °].
    """
    hydrogen_table = document.add_table(rows=len(data) + 1, cols=5, style='Table Grid')
    head_row = hydrogen_table.rows[0].cells
    # D-H...A	d(D-H)	d(H...A)	d(D...A)	<(DHA)
    head_row[0].paragraphs[0].add_run(
        f'D{halbgeviert}H{ellipsis_mid}A{protected_space}[{angstrom}]').font.bold = True
    head_row[1].paragraphs[0].add_run(f'd(D{halbgeviert}H){protected_space}[{angstrom}]').font.bold = True
    head_row[2].paragraphs[0].add_run(f'd(H{ellipsis_mid}A){protected_space}[{angstrom}]').font.bold = True
    head_row[3].paragraphs[0].add_run(f'd(D{ellipsis_mid}A){protected_space}[{angstrom}]').font.bold = True
    head_row[4].paragraphs[0].add_run(f'<(DHA){protected_space}[{degree_sign}]').font.bold = True
    col0_cells = hydrogen_table.columns[0].cells
    col1_cells = hydrogen_table.columns[1].cells
    col2_cells = hydrogen_table.columns[2].cells
    col3_cells = hydrogen_table.columns[3].cells
    col4_cells = hydrogen_table.columns[4].cells
    rowidx = 1
    for h in data.hydrogen_bonds_as_str:
        c0, c1, c2, c3, c4 = col0_cells[rowidx], col1_cells[rowidx], \
            col2_cells[rowidx], col3_cells[rowidx], col4_cells[rowidx]
        rowidx += 1
        c0.text = h.atoms
        c0.paragraphs[0].add_run(h.symm).font.superscript = True
        c1.text = h.dist_dh
        c2.text = h.dist_ha
        c3.text = h.dist_da
        c4.text = h.angle_dha
    widths = (Cm(4), Cm(2.5), Cm(2.5), Cm(2.5), Cm(2.5))
    make_table_widths(hydrogen_table, widths)
    if data.symmetry_generated_atoms_used:
        add_last_symminfo_line(data.symminfo, document)
    return table_num


def paragraph(table: Table) -> Paragraph:
    row_cells = table.add_row().cells
    p: Paragraph = row_cells[0].paragraphs[0]
    return p


def set_column_width(column, width: Length) -> None:
    for cell in column.cells:
        cell.width = width


def make_table_widths(table: Table, widths: Sequence[Length]) -> None:
    """
    Sets the width of the columns of a table.
    """
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def add_hydrogen_omit_info(document: Document) -> None:
    """
    Adds information if hydrogen atom bond are omitted from the respective table.
    """
    p = document.add_paragraph('')
    line = 'Bonds to hydrogen atoms were omitted.'
    run = p.add_run(line)
    run.font.size = Pt(8)


def add_last_symminfo_line(newsymms: str, document: Document) -> None:
    """
    Adds text about the symmetry generators used in order to add symmetry generated atoms.
    """
    p = document.add_paragraph('')
    run = p.add_run(newsymms)
    run.font.size = Pt(8)


def get_card(cif: CifContainer, symm: str) -> list[str]:
    """
    Returns a symmetry card from the _space_group_symop_operation_xyz or _symmetry_equiv_pos_as_xyz list.
    :param cif: the cif file object
    :param symm: the symmetry number
    :return: ['x', ' y', ' z'] etc
    """
    card = cif.symmops[int(symm.split('_')[0]) - 1].split(',')
    return card


if __name__ == '__main__':
    output_filename = 'tables.docx'
    import time

    settings = FinalCifSettings()
    options = Options(None, settings)

    # make_report_from(options, CifContainer('test-data/hydrogen/some_riding_some_isotropic.cif'),
    #                 output_filename='test.docx')
    make_report_from(options, CifContainer('test-data/DK_Zucker2_0m.cif'), output_filename='test.docx')
    # make_report_from(options, CifContainer(r'C:\Users\daniel.kratzert\Downloads\hydrogen_bond_types\1218_31_7_0m.cif'), output_filename='test.docx')
    # make_report_from(options, CifContainer(r'C:\Users\daniel.kratzert\Downloads\rqt_c1_0m_sq_complete.cif'), output_filename='test.docx')

    # make_multi_tables(CifContainer('test-data/1000007-multi-finalcif.cif'))
    open_file(Path('test.docx'))
    t1 = time.perf_counter()
