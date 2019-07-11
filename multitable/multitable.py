# -*- coding: utf-8 -*-
#
# This program was innitially created by Nils Trapp and extended by
# Daniel Kratzert
#
import itertools as it
import re
from pathlib import Path
from typing import List

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Pt, RGBColor


# compiled with "Py -3 -m PyInstaller multitable.spec --onefile"
from cif.cif_file_parser import Cif
from cif.file_reader import CifContainer
from multitable.tools import isfloat, grouper, cif_keywords_list, this_or_quest, get_files_from_current_dir


def populate_description_columns(table):
    """
    This Method adds the descriptions to the fist table column.
    """
    lgnd1 = table.cell(1, 0).paragraphs[0]
    lgnd1sub = lgnd1.add_run('Empirical formula')
    lgnd2 = table.cell(2, 0).paragraphs[0]
    lgnd2sub = lgnd2.add_run('Formula weight')
    lgnd3 = table.cell(3, 0).paragraphs[0]
    lgnd3sub = lgnd3.add_run('Temperature/K')
    lgnd4 = table.cell(4, 0).paragraphs[0]
    lgnd4sub = lgnd4.add_run('Crystal system')
    lgnd5 = table.cell(5, 0).paragraphs[0]
    lgnd5sub = lgnd5.add_run('Space group')
    lgnd6 = table.cell(6, 0).paragraphs[0]
    lgnd6sub = lgnd6.add_run('a').font.italic = True
    lgnd6.add_run('/\u212B')
    lgnd7 = table.cell(7, 0).paragraphs[0]
    lgnd7sub = lgnd7.add_run('b').font.italic = True
    lgnd7.add_run('/\u212B')
    lgnd8 = table.cell(8, 0).paragraphs[0]
    lgnd8sub = lgnd8.add_run('c').font.italic = True
    lgnd8.add_run('/\u212B')
    lgnd9 = table.cell(9, 0).paragraphs[0]
    lgnd9sub = lgnd9.add_run('\u03B1/\u00b0')
    lgnd10 = table.cell(10, 0).paragraphs[0]
    lgnd10sub = lgnd10.add_run('\u03B2/\u00b0')
    lgnd11 = table.cell(11, 0).paragraphs[0]
    lgnd11sub = lgnd11.add_run('\u03B3/\u00b0')
    lgnd12 = table.cell(12, 0).paragraphs[0]
    lgnd12sub = lgnd12.add_run('Volume/\u212B')
    lgnd12sub1 = lgnd12.add_run('3')
    lgnd12sub1.font.superscript = True
    lgnd13 = table.cell(13, 0).paragraphs[0]
    lgnd13sub = lgnd13.add_run('Z')
    lgnd13sub.font.italic = True
    lgnd14 = table.cell(14, 0).paragraphs[0]
    lgnd14sub = lgnd14.add_run('\u03C1')
    lgnd14sub.font.italic = True
    lgnd14sub1 = lgnd14.add_run('calc')
    lgnd14sub1.font.subscript = True
    lgnd14sub2 = lgnd14.add_run(' g/cm')
    lgnd14sub3 = lgnd14.add_run('3')
    lgnd14sub3.font.superscript = True
    lgnd15 = table.cell(15, 0).paragraphs[0]
    lgnd15sub = lgnd15.add_run('\u03BC/mm')
    lgnd15sub1 = lgnd15.add_run('-1')
    lgnd15sub1.font.superscript = True
    lgnd16 = table.cell(16, 0).paragraphs[0]
    lgnd16sub = lgnd16.add_run('F')
    lgnd16sub.font.italic = True
    lgnd16sub1 = lgnd16.add_run('(000)')
    lgnd17 = table.cell(17, 0).paragraphs[0]
    lgnd17sub = lgnd17.add_run('Crystal size/mm')
    lgnd17sub1 = lgnd17.add_run('3')
    lgnd17sub1.font.superscript = True
    lgnd18 = table.cell(18, 0).paragraphs[0]
    lgnd18sub = lgnd18.add_run('Crystal colour')
    lgnd19 = table.cell(19, 0).paragraphs[0]
    lgnd19sub = lgnd19.add_run('Crystal shape')
    lgnd20 = table.cell(20, 0).paragraphs[0]
    lgnd20sub = lgnd20.add_run('Radiation')
    lgnd21 = table.cell(21, 0).paragraphs[0]
    lgnd21sub = lgnd21.add_run('2\u03F4 range/\u00b0')
    lgnd22 = table.cell(22, 0).paragraphs[0]
    lgnd22sub = lgnd22.add_run('Index ranges')
    lgnd23 = table.cell(23, 0).paragraphs[0]
    lgnd23sub = lgnd23.add_run('Reflections collected')
    lgnd24 = table.cell(24, 0).paragraphs[0]
    lgnd24sub = lgnd24.add_run('Independent reflections')
    lgnd25 = table.cell(25, 0).paragraphs[0]
    lgnd25sub = lgnd25.add_run('Data / Restraints / Param.')
    lgnd26 = table.cell(26, 0).paragraphs[0]
    lgnd26sub = lgnd26.add_run('Goodness-of-fit on ')
    lgnd26sub1 = lgnd26.add_run('F')
    lgnd26sub1.font.italic = True
    lgnd26sub2 = lgnd26.add_run('2')
    lgnd26sub2.font.superscript = True
    lgnd27 = table.cell(27, 0).paragraphs[0]
    lgnd27sub = lgnd27.add_run('Final ')
    lgnd27sub1 = lgnd27.add_run('R')
    lgnd27sub1.font.italic = True
    lgnd27sub2 = lgnd27.add_run(' indexes \n[')
    lgnd27sub3 = lgnd27.add_run('I')
    lgnd27sub3.font.italic = True
    lgnd27sub4 = lgnd27.add_run('\u22652\u03C3(')
    lgnd27sub5 = lgnd27.add_run('I')
    lgnd27sub5.font.italic = True
    lgnd27sub3 = lgnd27.add_run(')]')
    lgnd28 = table.cell(28, 0).paragraphs[0]
    lgnd28sub = lgnd28.add_run('Final ')
    lgnd28sub1 = lgnd28.add_run('R')
    lgnd28sub1.font.italic = True
    lgnd28sub2 = lgnd28.add_run(' indexes \n[all data]')
    lgnd29 = table.cell(29, 0).paragraphs[0]
    lgnd29sub = lgnd29.add_run('Largest peak/hole /e\u212B')
    lgnd29sub1 = lgnd29.add_run('3')
    lgnd29sub1.font.superscript = True
    lgnd30 = table.cell(30, 0).paragraphs[0]
    lgnd30sub = lgnd30.add_run('Flack x parameter')


def format_space_group(table, cif, table_column):
    """
    Sets formating of the space group symbol in row 6.
    """
    # The HM space group symbol
    space_group = cif['_space_group_name_H-M_alt']
    it_number = cif['_space_group_IT_number']
    if space_group:
        if len(space_group) > 4:  # don't modify P 1
            space_group = re.sub(r'\s1', '', space_group)  # remove extra Hall "1" for mono and tric
        space_group = re.sub(r'\s', '', space_group)  # remove all remaining whitespace
        # space_group = re.sub(r'-1', u'\u0031\u0305', space_group)  # exchange -1 with 1bar
        space_group_formated_text = [char for char in space_group]  # ???)
        sgrun = table.cell(5, table_column + 1).paragraphs[0]
        is_sub = False
        for k in range(0, len(space_group_formated_text)):
            sgrunsub = sgrun.add_run(space_group_formated_text[k])
            if not space_group_formated_text[k].isdigit():
                sgrunsub.font.italic = True
            else:
                if space_group_formated_text[k - 1].isdigit() and not is_sub:
                    is_sub = True
                    sgrunsub.font.subscript = True  # lowercase the second digit if previous is also digit
                else:
                    is_sub = False  # only every second number as subscript for P212121 etc.
        if it_number:
            sgrun.add_run(' (' + it_number + ')')
    else:
        space_group = 'no space group'
    return space_group


def make_report_from(files: List, output_filename: str = None):
    """
    Creates a tabular cif report.
    :param files: Input cif files a list.
    :param output_filename: the table is saved to this file.
    """
    group_of_files = list(grouper(files, 3))  # group in threes to fit on A4 page
    table_index = len(group_of_files) - 1  # n-th table

    document = Document()

    style = document.styles['Normal']
    font = style.font
    font.name = 'Callibri'
    font.size = Pt(10)

    strHead = 'Structure Tables'
    # str1Par = '\nFormatting by user needed: font for symbols, table cell widths!'
    # str2Par = '\nPLEASE DOUBLE-CHECK ALL ENTRIES, CIF FILES CAN BE INCOMPLETE!!!'
    # str3Par = '\nParsed files:\n\n' + "\n".join(files)

    # a style for the header:
    styles = document.styles
    new_heading_style = styles.add_style('HeaderStyle', WD_STYLE_TYPE.PARAGRAPH)
    new_heading_style.base_style = styles['Heading 1']
    font = new_heading_style.font
    font.color.rgb = RGBColor(0, 0, 0)
    head = document.add_heading(strHead, 1)
    head.style = 'HeaderStyle'
    head.style.paragraph_format.space_before = Pt(0)
    # document.add_paragraph(str1Par)
    # document.add_paragraph(str2Par)
    # document.add_paragraph(str3Par)

    # document.add_page_break()

    for page_number, file_list in enumerate(group_of_files):  # one page per three structures:
        document.add_paragraph('')  # cannot format cells directly,
        paragraph = document.paragraphs[-1]  # but it will keep settings from
        paragraph_format = style.paragraph_format  # previous paragraph -> dirty hack:
        paragraph_format.space_before = Pt(2.5)  # create paragraph, apply style,
        paragraph_format.space_after = Pt(0)  # kill paragraph, create table.
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None
        table = document.add_table(rows=1, cols=4)
        header_cells = table.rows[0].cells
        table.autofit = False
        col = table.columns[0]
        col.width = Cm(4.5)
        # col.autofit = True
        col = table.columns[1]
        col.width = Cm(3.4)
        col = table.columns[2]
        col.width = Cm(3.4)
        col = table.columns[3]
        col.width = Cm(3.4)

        sum_formula = 'no sum formula'

        # setup table format:
        for cell in range(0, 30):
            row = table.add_row()  # define row and cells separately
            for table_column in range(0, 3):
                row.cells[table_column].style = document.styles['Normal']

        populate_description_columns(table)

        for table_column in range(0, 3):  # the three columns
            file_name = file_list[table_column]
            if not file_name:
                break
            file_obj = Path(file_name)
            if file_obj.exists():
                cif = CifContainer(file_obj)
                cif.open_cif_with_gemmi()

                headcell = header_cells[table_column].paragraphs[0]
                headcell.add_run(Path(file_list[table_column]).name).bold = True

                # Set text for all usual cif keywords:
                for _, key in enumerate(cif_keywords_list):
                    cell = table.cell(key[1] + 1, table_column + 1)
                    print(cif[key[0]])
                    if cif[key[0]]:
                        cell.text = cif[key[0]]
                    else:
                        cell.text = '?'
                        continue

                # Now the special handling:

                # The sum formula:
                if cif['_chemical_formula_sum']:
                    sum_formula = cif['_chemical_formula_sum']
                    ltext2 = sum_formula.replace(" ", "").replace("'", "")
                    ltext3 = [''.join(x[1]) for x in it.groupby(ltext2, lambda x: x.isalpha())]
                    for _, word in enumerate(ltext3):
                        formrun = table.cell(1, table_column + 1).paragraphs[0]
                        formrunsub = formrun.add_run(word)
                        if isfloat(word):
                            formrunsub.font.subscript = True

                space_group = format_space_group(table, cif, table_column)
                radiation_type = cif['_diffrn_radiation_type']
                radiation_wavelength = cif['_diffrn_radiation_wavelength']
                crystal_size_min = cif['_exptl_crystal_size_min']
                crystal_size_mid = cif['_exptl_crystal_size_mid']
                crystal_size_max = cif['_exptl_crystal_size_max']
                limit_h_min = cif['_diffrn_reflns_limit_h_min']
                limit_h_max = cif['_diffrn_reflns_limit_h_max']
                limit_k_min = cif['_diffrn_reflns_limit_k_min']
                limit_k_max = cif['_diffrn_reflns_limit_k_max']
                theta_min = cif['_diffrn_reflns_theta_min']
                theta_max = cif['_diffrn_reflns_theta_max']
                limit_l_min = cif['_diffrn_reflns_limit_l_min']
                limit_l_max = cif['_diffrn_reflns_limit_l_max']
                reflns_number_total = cif['_reflns_number_total']
                reflns_av_R_equivalents = cif['_diffrn_reflns_av_R_equivalents']
                reflns_av_unetI = cif['_diffrn_reflns_av_unetI/netI']
                ls_number_reflns = cif['_refine_ls_number_reflns']
                ls_number_restraints = cif['_refine_ls_number_restraints']
                ls_number_parameters = cif['_refine_ls_number_parameters']
                ls_R_factor_gt = cif['_refine_ls_R_factor_gt']
                ls_wR_factor_gt = cif['_refine_ls_wR_factor_gt']
                ls_R_factor_all = cif['_refine_ls_R_factor_all']
                ls_wR_factor_ref = cif['_refine_ls_wR_factor_ref']
                try:
                    diff_density_min = "{0:.2f}".format(round(float(cif['_refine_diff_density_min']), 2))
                except ValueError:
                    diff_density_min = '?'
                try:
                    diff_density_max = "{0:.2f}".format(round(float(cif['_refine_diff_density_max']), 2))
                except ValueError:
                    diff_density_max = '?'

                # now prepare & write all the concatenated & derived cell contents:
                table.cell(17, table_column + 1).text = this_or_quest(crystal_size_max) + '\u00d7' + \
                                                        this_or_quest(crystal_size_mid) + '\u00d7' + \
                                                        this_or_quest(crystal_size_min)
                wavelength = str(' (\u03bb =' + this_or_quest(radiation_wavelength) + ')').replace(' ', '')
                # radtype: ('Mo', 'K', '\\a')
                radtype = list(radiation_type.partition("K"))
                if len(radtype) > 2:
                    if radtype[2] == '\\a':
                        radtype[2] = '\u03b1'
                    if radtype[2] == '\\b':
                        radtype[2] = '\u03b2'
                radrun = table.cell(20, table_column + 1).paragraphs[0]
                # radiation type e.g. Mo:
                radrun.add_run(radtype[0])
                # K line:
                radrunita = radrun.add_run(radtype[1])
                radrunita.font.italic = True
                alpha = radrun.add_run(radtype[2])
                alpha.font.italic = True
                alpha.font.subscript = True
                # wavelength lambda:
                radrun.add_run(' ' + wavelength)
                try:
                    table.cell(21, table_column + 1).text = "{0:.2f}".format(2 * float(theta_min)) + \
                                                            ' to ' + "{0:.2f}".format(2 * float(theta_max))
                except ValueError:
                    table.cell(21, table_column + 1).text = '? to ?'
                table.cell(22, table_column + 1).text = limit_h_min + ' \u2264 h \u2264 ' \
                                                        + limit_h_max + '\n' \
                                                        + limit_k_min + ' \u2264 k \u2264 ' \
                                                        + limit_k_max + '\n' \
                                                        + limit_l_min + ' \u2264 l \u2264 ' \
                                                        + limit_l_max
                rintrun = table.cell(24, table_column + 1).paragraphs[0]
                rintrun.add_run(this_or_quest(reflns_number_total) + '\n')
                rintita1 = rintrun.add_run('R')
                rintita1.font.italic = True
                rintsub1 = rintrun.add_run('int')
                rintsub1.font.subscript = True
                rintrun.add_run(' = ' + this_or_quest(reflns_av_R_equivalents) + '\n')
                rintita2 = rintrun.add_run('R')
                rintita2.font.italic = True
                rintsub2 = rintrun.add_run('sigma')
                rintsub2.font.subscript = True
                rintrun.add_run(' = ' + this_or_quest(reflns_av_unetI))
                table.cell(25, table_column + 1).text = this_or_quest(ls_number_reflns) + '/' \
                                                        + this_or_quest(ls_number_restraints) + '/' \
                                                        + this_or_quest(ls_number_parameters)
                r2sigrun = table.cell(27, table_column + 1).paragraphs[0]
                r2sigita1 = r2sigrun.add_run('R')
                r2sigita1.font.italic = True
                r2sigsub1 = r2sigrun.add_run('1')
                r2sigsub1.font.subscript = True
                r2sigrun.add_run(' = ' + this_or_quest(ls_R_factor_gt) + '\nw')
                r2sigita2 = r2sigrun.add_run('R')
                r2sigita2.font.italic = True
                r2sigsub2 = r2sigrun.add_run('2')
                r2sigsub2.font.subscript = True
                r2sigrun.add_run(' = ' + this_or_quest(ls_wR_factor_gt))
                rfullrun = table.cell(28, table_column + 1).paragraphs[0]
                rfullita1 = rfullrun.add_run('R')
                rfullita1.font.italic = True
                rfullsub1 = rfullrun.add_run('1')
                rfullsub1.font.subscript = True
                rfullrun.add_run(' = ' + this_or_quest(ls_R_factor_all) + '\nw')
                rfullita2 = rfullrun.add_run('R')
                rfullita2.font.italic = True
                rfullsub2 = rfullrun.add_run('2')
                rfullsub2.font.subscript = True
                rfullrun.add_run(' = ' + ls_wR_factor_ref)
                table.cell(29, table_column + 1).text = diff_density_max + '/' + diff_density_min
                table.cell(30, table_column + 1).text = cif['_refine_ls_abs_structure_Flack']
                print('File parsed: ' + file_obj.name + '  (' + sum_formula + ')  ' + space_group)

        # page break between tables:
        if page_number < table_index:
            document.add_page_break()

    print('\nScript finished - output file: multitable.docx')
    if not output_filename:
        document.save('multitable.docx')
    else:
        document.save(output_filename)
    return group_of_files


if __name__ == '__main__':
    make_report_from(get_files_from_current_dir())
