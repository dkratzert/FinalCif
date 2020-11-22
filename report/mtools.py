import re

from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.text.paragraph import Paragraph

from report.report_text import math_to_word
from report.spgrps import SpaceGroups

cif_keywords_list = (
    ['_chemical_formula_weight', 1],
    ['_diffrn_ambient_temperature', 2],
    ['_space_group_crystal_system', 3],
    ['_cell_length_a', 5],
    ['_cell_length_b', 6],
    ['_cell_length_c', 7],
    ['_cell_angle_alpha', 8],
    ['_cell_angle_beta', 9],
    ['_cell_angle_gamma', 10],
    ['_cell_volume', 11],
    ['_cell_formula_units_Z', 12],
    ['_exptl_crystal_density_diffrn', 13],
    ['_exptl_absorpt_coefficient_mu', 14],
    ['_exptl_crystal_F_000', 15],
    ['_exptl_crystal_colour', 17],
    ['_exptl_crystal_description', 18],
    ['_diffrn_reflns_theta_min', 20],
    ['_diffrn_reflns_theta_max', 20],
    ['_diffrn_reflns_number', 22],
    ['_refine_ls_goodness_of_fit_ref', 25],
)


def format_space_group(paragraph: Paragraph, space_group: str, it_number) -> None:
    """
    Sets formating of the space group symbol in row 6.
    """
    try:
        # The HM space group symbol
        s = SpaceGroups()
        spgrxml = s.iucrNumberToMathml(it_number)
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        paragraph._element.append(math_to_word(spgrxml))
        paragraph.add_run(' (' + it_number + ')')
    except Exception:
        # Use fallback:
        if space_group:
            if len(space_group) > 4:  # don't modify P 1
                space_group = re.sub(r'\s1', '', space_group)  # remove extra Hall "1" for mono and tric
            space_group = re.sub(r'\s', '', space_group)  # remove all remaining whitespace
            # space_group = re.sub(r'-1', one_bar, space_group)  # exchange -1 with 1bar
            space_group_formated_text = [char for char in space_group]  # ???)
            is_sub = False
            for k, char in enumerate(space_group_formated_text):
                sgrunsub = paragraph.add_run(char)
                if not char.isdigit():
                    sgrunsub.font.italic = True
                else:
                    if space_group_formated_text[k - 1].isdigit() and not is_sub:
                        is_sub = True
                        sgrunsub.font.subscript = True  # lowercase the second digit if previous is also digit
                    else:
                        is_sub = False  # only every second number as subscript for P212121 etc.
            if it_number:
                paragraph.add_run(' (' + it_number + ')')
        else:
            paragraph.add_run('?')
