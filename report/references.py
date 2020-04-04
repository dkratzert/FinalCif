#  ----------------------------------------------------------------------------
#  "THE BEER-WARE LICENSE" (Revision 42):
#  daniel.kratzert@ac.uni-freiburg.de> wrote this file.  As long as you retain
#  this notice you can do whatever you want with this stuff. If we meet some day,
#  and you think this stuff is worth it, you can buy me a beer in return.
#  Dr. Daniel Kratzert
#  ----------------------------------------------------------------------------
from contextlib import suppress
from typing import List

from docx.text.paragraph import Paragraph

"""
[1] SAINT
[2] SADABS/TWINABS
[3] SHELXT
[4] SHELXL
[5] ShelXle
[6] DSR
[7] Olex2
[8] Absolute structure (nach S. Parson)
[9] Überstrukturen (A non-mathematical introduction to superstructures)

[1] SAINT V8.37A, Bruker AXS, Madison, Wisconsin, USA, 2015.
[2] L. Krause, R. Herbst-Irmer, G. M. Sheldrick, D. Stalke, J. Appl. Cryst. 2015, 48, 3–10, doi:10.1107/S1600576714022985
[3] G. M. Sheldrick, Acta Cryst. 2015, A71, 3–8, doi:10.1107/S2053273314026370
[4] G. M. Sheldrick, Acta Cryst. 2015, C71, 3–8, doi:10.1107/S2053229614024218.
[5] C. B. Hubschle, G. M. Sheldrick and B. Dittrich, J. Appl. Cryst. 2011, 44, 1281–1284, 10.1107/S0021889811043202
[6a] D. Kratzert, I. Krossing, J. Appl. Cryst. 2018, 51, 928-934. doi: 10.1107/S1600576718004508
[6b] D. Kratzert, J. J. Holstein, I. Krossing, J. Appl. Cryst. 2015, 48, 933–938, doi:10.1107/S1600576715005580
[7] L. J. Bourhis, O. V. Dolomanov, R. J. Gildea, J. A. K. Howard, H. Puschmann, Acta Cryst. 2015, A71, 59–75, doi:10.1107/S2053273314022207
[8] S. Parsons, H. D. Flack, T. Wagner, Acta Cryst. 2013, B69, 249-259, doi:10.1107/S2052519213010014
[9] T. Wagner, A. Schonleber, Acta Cryst. 2009, 65, 249–268, doi:10.1107/S0108768109015614

Missing: d*trek, SORTAV, Crysalispro, CCDC, 
"""


class ReferenceList():
    """
    This reference list holds a list of all used references. During each self.append(Reference()), a new reference is
    appended to the list. 
    At the end of the document, a numbered list of references can be generated with self.make_make_literature_list().
    """

    def __init__(self, paragraph: Paragraph):
        self.paragraph = paragraph
        self._references = list()

    def append(self, ref) -> None:
        """Adds a superscript list of reference numbers in brackets to the document."""
        if isinstance(ref, (list, tuple)):
            if not ref:
                return None
            self._append_list(ref)
        else:
            if not ref in self._references:
                self._references.append(ref)
            self.paragraph.add_run('[{}]'.format(self._references.index(ref) + 1)).font.superscript = True
        # better not here:
        # self.paragraph.add_run(' ')

    def _append_list(self, reflist: List) -> None:
        reflst_long = []
        self.paragraph.add_run('[').font.superscript = True
        for n, ref in enumerate(reflist):
            if not ref in self._references:
                self._references.append(ref)
            refnum = self._references.index(ref) + 1
            reflst_long.append(refnum)
        reftxt = self.get_sequence(reflst_long)
        self.paragraph.add_run(reftxt).font.superscript = True
        self.paragraph.add_run(']').font.superscript = True

    @staticmethod
    def get_sequence(stringlist: List[int]):
        """
        Converts a list of numbers into a string of numbers where recurring sequences
        are described with a range.

        >>> ReferenceList.get_sequence([1, 3, 4, 5, 6, 8, 11])
        '1,3-6,8,11'
        """
        folg = []
        start = 0
        for n, val in enumerate(stringlist):
            val = int(val)
            nextval = 0
            with suppress(IndexError):
                nextval = int(stringlist[n + 1])
            nextnext = 0
            with suppress(IndexError):
                nextnext = int(stringlist[n + 2])
            # a sequence starts:
            if nextnext == val + 2 and not start:
                start = val
            # in a sequence and next value is not +1 -> squence ends with val:
            if start and nextval != val + 1:
                folg.append('{}-{}'.format(start, val))
                start = 0
                continue
            # everything outside a sequence:
            if not start:
                folg.append(val)
        return ','.join([str(x) for x in folg])

    def make_literature_list(self, paragraph_reflist):
        for num, ref in enumerate(self._references, 1):
            paragraph_reflist.add_run('[{}] '.format(str(num)))
            ref.add_reference(paragraph_reflist)
            paragraph_reflist.add_run('\n')


class ReferenceFormater():
    def __init__(self):
        self.authors = ''
        self.journal = ''
        self.year = ''
        self.volume = ''
        self.pages = ''
        self.doi = ''
        self.program = ''

    def add_reference(self, p):
        if self.authors:
            p.add_run(self.authors)
            p.add_run(', ')
        if self.journal:
            p.add_run(self.journal).italic = True
            if not self.journal.endswith('.'):
                p.add_run(', ')
            else:
                p.add_run(' ')
        if self.year:
            p.add_run(self.year).bold = True
            p.add_run(', ')
        if self.volume:
            p.add_run(self.volume).italic = True
            p.add_run(', ')
        if self.pages:
            p.add_run(self.pages)
            if self.doi:
                p.add_run(', ')
        if self.doi:
            p.add_run(self.doi)
        if any([self.journal, self.pages, self.year, self.volume, self.doi]):
            p.add_run('.')

    def __repr__(self):
        txt = ''
        if self.authors:
            txt += self.authors
            txt += ', '
        if self.journal:
            txt += self.journal
            if not self.journal.endswith('.'):
                txt += ', '
            else:
                txt += ' '
        if self.year:
            txt += self.year
            txt += ', '
        if self.volume:
            txt += self.volume
            txt += ', '
        if self.pages:
            txt += self.pages
            if self.doi:
                txt += ', '
        if self.doi:
            txt += self.doi
        if any([self.journal, self.pages, self.year, self.volume, self.doi]):
            txt += '.'
        return txt


class DummyReference(ReferenceFormater):
    """
    >>> DummyReference()
    Unknown Reference, please add.
    """

    def __init__(self):
        super().__init__()
        self.doi = 'Unknown Reference, please add'


class DSRReference2015(ReferenceFormater):
    """
    >>> DSRReference2015()
    D. Kratzert, J.J. Holstein, I. Krossing, J. Appl. Cryst. 2015, 48, 933–938, doi:10.1107/S1600576715005580.
    """

    def __init__(self):
        super().__init__()
        # self.doi = '(doi: 10.1107/S1600576715005580)'
        self.authors = 'D. Kratzert, J.J. Holstein, I. Krossing'
        self.journal = 'J. Appl. Cryst.'
        self.year = '2015'
        self.volume = '48'
        self.pages = '933–938'
        self.doi = 'doi:10.1107/S1600576715005580'


class DSRReference2018(ReferenceFormater):
    """
    >>> DSRReference2018()
    D. Kratzert, I. Krossing, J. Appl. Cryst. 2018, 51, 928–934, doi:10.1107/S1600576718004508.
    """

    def __init__(self):
        super().__init__()
        self.doi = 'doi:10.1107/S1600576718004508'
        self.authors = 'D. Kratzert, I. Krossing'
        self.journal = 'J. Appl. Cryst.'
        self.year = '2018'
        self.volume = '51'
        self.pages = '928–934'


class BrukerReference(ReferenceFormater):
    def __init__(self, name: str, version: str):
        """
        >>> BrukerReference('SAINT', 'V7.68a')
        Bruker, SAINT, V7.68a, Bruker AXS Inc., Madison, Wisconsin, USA.
        """
        super().__init__()
        self.authors = 'Bruker'
        self.journal = name
        # self.year = '2012'
        # if '6.28' in self.year:
        #    self.year = '2001'
        self.volume = version
        self.pages = 'Bruker AXS Inc., Madison, Wisconsin, USA'


class SADABS_TWINABS_Reference(ReferenceFormater):
    def __init__(self):
        """
        L. Krause, R. Herbst-Irmer, G. M. Sheldrick, D. Stalke, J. Appl. Cryst. 2015, 48, 3–10,
            doi:10.1107/S1600576714022985

        >>> SADABS_TWINABS_Reference()
        L. Krause, R. Herbst-Irmer, G. M. Sheldrick, D. Stalke, J. Appl. Cryst. 2015, 48, 3–10, doi:10.1107/S1600576714022985.
        """
        super().__init__()
        self.authors = 'L. Krause, R. Herbst-Irmer, G. M. Sheldrick, D. Stalke'
        self.journal = 'J. Appl. Cryst.'
        self.year = '2015'
        self.volume = '48'
        self.pages = '3–10'
        self.doi = 'doi:10.1107/S1600576714022985'


class SHELXTReference(ReferenceFormater):
    def __init__(self):
        """
        >>> SHELXTReference()
        G. M. Sheldrick, Acta Cryst. 2015, A71, 3–8, doi:10.1107/S2053273314026370.
        """
        super().__init__()
        self.authors = 'G. M. Sheldrick'
        self.journal = 'Acta Cryst.'
        self.year = '2015'
        self.volume = 'A71'
        self.pages = '3–8'
        self.doi = 'doi:10.1107/S2053273314026370'


class SHELXSReference(ReferenceFormater):
    def __init__(self):
        """
        >>> SHELXSReference()
        G. M. Sheldrick, Acta Cryst. 2008, A64, 112–122, doi:10.1107/S0108767307043930.
        """
        super().__init__()
        self.authors = 'G. M. Sheldrick'
        self.journal = 'Acta Cryst.'
        self.year = '2008'
        self.volume = 'A64'
        self.pages = '112–122'
        self.doi = 'doi:10.1107/S0108767307043930'


class SHELXLReference(ReferenceFormater):
    """
    >>> SHELXLReference()
    G. M. Sheldrick, Acta Cryst. 2015, C71, 3–8, doi:10.1107/S2053229614024218.
    """

    def __init__(self):
        super(SHELXLReference, self).__init__()
        self.authors = 'G. M. Sheldrick'
        self.year = '2015'
        self.journal = 'Acta Cryst.'
        self.volume = 'C71'
        self.pages = '3–8'
        self.doi = 'doi:10.1107/S2053229614024218'


class ShelXleReference(ReferenceFormater):
    """
    >>> ShelXleReference()
    C. B. Hubschle, G. M. Sheldrick, B. Dittrich, J. Appl. Cryst. 2011, 44, 1281–1284, doi:10.1107/S0021889811043202.
    """

    def __init__(self):
        super().__init__()
        self.authors = 'C. B. Hubschle, G. M. Sheldrick, B. Dittrich'
        self.year = '2011'
        self.journal = 'J. Appl. Cryst.'
        self.volume = '44'
        self.pages = '1281–1284'
        self.doi = 'doi:10.1107/S0021889811043202'


class Olex2Reference(ReferenceFormater):
    """
    >>> Olex2Reference()
    L. J. Bourhis, O. V. Dolomanov, R. J. Gildea, J. A. K. Howard, H. Puschmann, Acta Cryst. 2015, A71, 59–75, doi:10.1107/S2053273314022207.
    """

    def __init__(self):
        super().__init__()
        self.authors = 'L. J. Bourhis, O. V. Dolomanov, R. J. Gildea, J. A. K. Howard, H. Puschmann'
        self.year = '2015'
        self.journal = 'Acta Cryst.'
        self.volume = 'A71'
        self.pages = '59–75'
        self.doi = 'doi:10.1107/S2053273314022207'


class ParsonFlackReference(ReferenceFormater):
    """
    >>> ParsonFlackReference()
    S. Parsons, H. D. Flack, T. Wagner, Acta Cryst. 2013, B69, 249–259, doi:10.1107/S2052519213010014.
    """

    def __init__(self):
        super().__init__()
        self.authors = 'S. Parsons, H. D. Flack, T. Wagner'
        self.year = '2013'
        self.journal = 'Acta Cryst.'
        self.volume = 'B69'
        self.pages = '249–259'
        self.doi = 'doi:10.1107/S2052519213010014'


class CCDCReference(ReferenceFormater):
    """
    >>> CCDCReference()
    C. R. Groom, I. J. Bruno, M. P. Lightfoot, S. C. Ward, Acta Cryst. 2016, B72, 171–179, doi:10.1107/S2052520616003954.
    """

    def __init__(self):
        super().__init__()
        self.authors = 'C. R. Groom, I. J. Bruno, M. P. Lightfoot, S. C. Ward'
        self.year = '2016'
        self.journal = 'Acta Cryst.'
        self.volume = 'B72'
        self.pages = '171–179'
        self.doi = 'doi:10.1107/S2052520616003954'


class SORTAVReference(ReferenceFormater):
    """
    >>> SORTAVReference()
    Robert H. Blessing, Cryst. Rev. 1987, 1, 3-58, doi:10.1080/08893118708081678.
    """

    def __init__(self):
        super().__init__()
        self.authors = 'Robert H. Blessing'
        self.year = '1987'
        self.journal = 'Cryst. Rev.'
        self.volume = '1'
        self.pages = '3-58'
        self.doi = 'doi:10.1080/08893118708081678'


class FinalCifReference(ReferenceFormater):
    """
    >>> FinalCifReference()
    D. Kratzert, FinalCif, 2020, https://www.xs3.uni-freiburg.de/research/finalcif.
    """

    def __init__(self):
        super().__init__()
        self.authors = 'D. Kratzert'
        self.year = '2020'
        self.journal = 'FinalCif'
        self.doi = 'https://www.xs3.uni-freiburg.de/research/finalcif'


if __name__ == '__main__':
    r = ReferenceList.get_sequence([1, 2, 3, 4])
    print(r)
