import os
from builtins import str

import gemmi
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from lxml import etree

from app_path import application_path
from cif.cif_file_io import CifContainer, retranslate_delimiter
from report.references import DummyReference, BrukerReference, SORTAVReference, ReferenceList, CCDCReference, \
    SHELXLReference, SHELXTReference, SHELXSReference, FinalCifReference, ShelXleReference, Olex2Reference
from tools.misc import prot_space, angstrom


def math_to_word(eq: str) -> str:
    """Transform a sympy equation to be printed in word document."""
    tree = etree.fromstring(eq)
    xslt = etree.parse(os.path.join(application_path, 'template/mathml2omml.xsl'))
    transform = etree.XSLT(xslt)
    new_dom = transform(tree)
    return new_dom.getroot()


def clean_string(string: str) -> str:
    """
    Removes control characters from a string.
    >>> clean_string('This is a sentence\\r with newline.')
    'This is a sentence with newline'
    >>> clean_string('')
    ''
    >>> clean_string('  This is  a sentence\\r with. newline.  ')
    'This is  a sentence with. newline'
    """
    return string \
        .replace('\n', '') \
        .replace('\r', '') \
        .replace('\t', '') \
        .replace('\f', '') \
        .replace('\0', '') \
        .strip(' ') \
        .strip('.')


def gstr(string: str) -> str:
    """
    Turn a string into a gemmi string and remove control characters.
    """
    return clean_string(gemmi.cif.as_string(string).strip("'"))


class FormatMixin():

    def bold(self, run: Run):
        r = run.bold = True
        return r


class ReportText():
    def __init__(self):
        pass


class Crystallization(FormatMixin):
    def __init__(self, cif: CifContainer, paragraph: Paragraph):
        self.cif = cif
        self.crytsalization_method = gstr(self.cif['_exptl_crystal_recrystallization_method']) + '.'
        if not self.crytsalization_method:
            self.crytsalization_method = '[No crystallization method was given]'
        sentence = "{} "
        self.text = sentence.format(self.crytsalization_method)
        paragraph.add_run(retranslate_delimiter(self.text))


class CrstalSelection(FormatMixin):
    def __init__(self, cif: CifContainer, paragraph: Paragraph):
        self.cif = cif
        self.temperature = gstr(self.cif['_diffrn_ambient_temperature'])
        self._name = cif.fileobj.name
        method = 'shock-cooled '
        sentence = "The data for {} were collected from a {}single crystal at {}{}K "
        try:
            if float(self.temperature.split('(')[0]) > 200:
                method = ''
        except ValueError:
            method = ''
        self.txt = sentence.format(self.name, method, self.temperature, prot_space)
        paragraph.add_run(retranslate_delimiter(self.txt))

    @property
    def name(self) -> str:
        if self._name:
            return self._name
        else:
            return '?'


class MachineType():
    def __init__(self, cif: CifContainer, paragraph: Paragraph):
        self.cif = cif
        self.difftype = gstr(self.cif['_diffrn_measurement_device_type']) \
                        or '[No measurement device type given]'
        self.device = gstr(self.cif['_diffrn_measurement_device']) \
                      or '[No measurement device given]'
        self.source = gstr(self.cif['_diffrn_source']).strip('\n\r') \
                      or '[No radiation source given]'
        self.monochrom = gstr(self.cif['_diffrn_radiation_monochromator']) \
                         or '[No monochromator type given]'
        if not self.monochrom:
            self.monochrom = '?'
        self.cooling = gstr(self.cif['_olex2_diffrn_ambient_temperature_device']) \
                       or ''
        self.rad_type = gstr(self.cif['_diffrn_radiation_type']) \
                        or '[No radiation type given]'
        radtype = format_radiation(self.rad_type)
        self.wavelen = gstr(self.cif['_diffrn_radiation_wavelength']) \
                       or '[No wavelength given]'
        self.detector_type = ''
        detector_type = gstr(self.cif['_diffrn_detector_type']) \
                        or '[No detector type given]'
        if detector_type:
            self.detector_type = " and a {} detector".format(detector_type)
        sentence1 = "on {} {} {} with {} {} using {} as monochromator{}. " \
                    "The diffractometer was equipped with {} {} low temperature device and used "
        sentence2 = " radiation (λ = {}" + prot_space + "{}). ".format(angstrom)
        txt = sentence1.format(get_inf_article(self.difftype), self.difftype, self.device,
                               get_inf_article(self.source), self.source, self.monochrom,
                               self.detector_type, get_inf_article(self.cooling), self.cooling)
        paragraph.add_run(retranslate_delimiter(txt))
        # radiation type e.g. Mo:
        paragraph.add_run(retranslate_delimiter(radtype[0]))
        # K line:
        radrunita = paragraph.add_run(radtype[1])
        radrunita.font.italic = True
        alpha = paragraph.add_run(retranslate_delimiter(radtype[2]))
        alpha.font.italic = True
        alpha.font.subscript = True
        txt2 = sentence2.format(self.wavelen)
        paragraph.add_run(txt2)


class DataReduct():
    def __init__(self, cif: CifContainer, paragraph: Paragraph, ref: ReferenceList):
        self.cif = cif
        integration = gstr(self.cif['_computing_data_reduction']) or '??'
        abstype = gstr(self.cif['_exptl_absorpt_correction_type']) or '??'
        abs_details = gstr(self.cif['_exptl_absorpt_process_details']) or '??'
        data_reduct_ref = DummyReference()
        absorpt_ref = DummyReference()
        integration_prog = '[unknown integration program]'
        scale_prog = ['unknown program']
        if 'SAINT' in integration:
            saintversion = 'unknown version'
            if len(integration.split()) > 0:
                saintversion = integration.split()[1]
            integration_prog = 'SAINT'
            data_reduct_ref = BrukerReference('SAINT', saintversion)
        absdetails = cif['_exptl_absorpt_process_details'].replace('-', ' ')
        if 'SADABS' in absdetails.upper() or 'TWINABS' in absdetails.upper():
            if len(absdetails.split()) > 0:
                version = absdetails.split()[1]
            else:
                version = 'unknown version'
            if 'SADABS' in absdetails:
                scale_prog = 'SADABS'
            else:
                scale_prog = 'TWINABS'
            absorpt_ref = BrukerReference(scale_prog, version)
        if 'SORTAV' in absdetails.upper():
            scale_prog = 'SORTAV'
            absorpt_ref = SORTAVReference()
        if 'crysalis' in abs_details.lower():
            scale_prog = 'SCALE3 ABSPACK'
        sentence = 'All data were integrated with {} and {} {} absorption correction using {} was applied.'
        txt = sentence.format(integration_prog,
                              get_inf_article(abstype),
                              abstype,
                              scale_prog)
        paragraph.add_run(retranslate_delimiter(txt))
        ref.append([data_reduct_ref, absorpt_ref])


class SolveRefine():
    def __init__(self, cif: CifContainer, paragraph: Paragraph, ref: ReferenceList):
        self.cif = cif
        refineref = DummyReference()
        solveref = DummyReference()
        solution_prog = gstr(self.cif['_computing_structure_solution']) or '??'
        solution_method = gstr(self.cif['_atom_sites_solution_primary']) or '??'
        if 'SHELXT' in solution_prog:
            solveref = SHELXTReference()
        if 'SHELXS' in solution_prog:
            solveref = SHELXSReference()
        refined = gstr(self.cif['_computing_structure_refinement']) or '??'
        if 'SHELXL' in refined.upper():
            refineref = SHELXLReference()
        if 'OLEX' in refined.upper():
            refineref = Olex2Reference()
        refine_coef = gstr(self.cif['_refine_ls_structure_factor_coef'])
        sentence = r"The structure were solved by {} methods using {} and refined by full-matrix " \
                   "least-squares methods against "
        txt = sentence.format(solution_method.strip('\n\r'), solution_prog.split()[0])
        paragraph.add_run(retranslate_delimiter(txt))
        paragraph.add_run('F').font.italic = True
        if refine_coef.lower() == 'fsqd':
            paragraph.add_run('2').font.superscript = True
        paragraph.add_run(' by {}'.format(refined.split()[0]))
        shelxle = None
        if 'shelxle' in refined.lower() or 'shelxle' in self.cif['_computing_molecular_graphics'].lower():
            paragraph.add_run(' using ShelXle')
            shelxle = ShelXleReference()
        paragraph.add_run('.')
        ref.append([solveref, refineref, shelxle])


class Hydrogens():
    def __init__(self, cif: CifContainer, paragraph: Paragraph):
        """
        TODO: check if the proposed things are really there.
        """
        self.cif = cif
        sentence1 = "All non-hydrogen atoms were refined with anisotropic displacement parameters. " \
                    "The hydrogen atoms were refined isotropically on calculated positions using a riding model " \
                    "with their "
        sentence2 = " values constrained to 1.5 times the "
        sentence3 = " of their pivot atoms for terminal sp"
        sentence4 = " carbon atoms and 1.2 times for all other carbon atoms."
        paragraph.add_run(sentence1)
        paragraph.add_run('U').font.italic = True
        paragraph.add_run('iso').font.subscript = True
        paragraph.add_run(sentence2)
        paragraph.add_run('U').font.italic = True
        paragraph.add_run('eq').font.subscript = True
        paragraph.add_run(sentence3)
        paragraph.add_run('3').font.superscript = True
        paragraph.add_run(sentence4)


class Disorder():
    def __init__(self, cif: CifContainer, paragraph: Paragraph):
        self.cif = cif
        self.dsr_sentence = ''
        sentence1 = "Disordered moieties were refined using bond lengths " \
                    "restraints and displacement parameter restraints. "
        if self.cif.dsr_used:
            self.dsr_sentence = "Some parts of the disorder model were introduced by the " \
                                "program DSR."
        paragraph.add_run(sentence1)
        if self.dsr_sentence:
            paragraph.add_run(self.dsr_sentence)


class SpaceChar(object):
    def __init__(self, paragraph: Paragraph):
        self.p = paragraph

    def regular(self):
        self.p.add_run(' ')

    def porotected(self):
        self.p.add_run(prot_space)


class CCDC():
    def __init__(self, cif: CifContainer, paragraph: Paragraph, ref: ReferenceList):
        self.cif = cif
        ccdc_num = gstr(self.cif['_database_code_depnum_ccdc_archive']) or '??????'
        sentence1 = "Crystallographic data (including structure factors) for the structures reported in this " \
                    "paper have been deposited with the Cambridge Crystallographic Data Centre."
        sentence2 = "CCDC {} contain " \
                    "the supplementary crystallographic data for this paper. Copies of the data can " \
                    "be obtained free of charge from the Cambridge Crystallographic Data Centre " \
                    "via www.ccdc.cam.ac.uk/structures.".format(ccdc_num)
        paragraph.add_run(sentence1)
        ref.append(CCDCReference())
        SpaceChar(paragraph).regular()
        paragraph.add_run(sentence2)


class FinalCifreport():
    def __init__(self, paragraph: Paragraph, ref: ReferenceList):
        sentence = "This report and the CIF file were generated using FinalCif."
        paragraph.add_run(sentence)
        ref.append(FinalCifReference())


def get_inf_article(next_word: str) -> str:
    if not next_word:
        return 'a'
    voc = 'aeiou'
    return 'a' if not next_word[0].lower() in voc else 'an'


def format_radiation(radiation_type: str) -> list:
    radtype = list(radiation_type.partition("K"))
    if len(radtype) > 2:
        radtype[2] = retranslate_delimiter(radtype[2])
        return radtype
    else:
        return radtype
