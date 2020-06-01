# -*- coding: utf-8 -*-
#
# This program was innitially created by Nils Trapp and extended by
# Daniel Kratzert
#
import itertools as it
import os
import subprocess
from math import sin, radians
from pathlib import Path
from typing import List, Sequence, Tuple

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Length, Pt
from docx.table import Table, _Cell

from app_path import application_path
from cif.cif_file_io import CifContainer
from report.mtools import cif_keywords_list, isfloat, this_or_quest, format_space_group
from report.references import ReferenceList, DSRReference2018, DSRReference2015
from report.report_text import CCDC, CrstalSelection, Crystallization, DataReduct, Disorder, Hydrogens, MachineType, \
    SolveRefine, format_radiation, FinalCifreport, SpaceChar
from report.symm import SymmetryElement
from tools.misc import prot_space, angstrom, bequal, sigma_sm, halbgeviert, degree_sign, ellipsis_mid, lessequal, \
    timessym, lambdasym


def make_report_from(options: dict, file_obj: Path, output_filename: str = None, path: str = '',
                     picfile: Path = None) -> str:
    """
    Creates a tabular cif report.
    :param file_obj: Input cif file.
    :param output_filename: the table is saved to this file.
    """
    without_H = options.get('without_H')
    report_text = options.get('report_text')
    document = create_document(path)

    if file_obj.exists():
        try:
            cif = CifContainer(file_obj)
        except Exception as e:
            print('Unable to open cif file:')
            print(e)
            raise
    else:
        raise FileNotFoundError
    if not cif:
        print('Something failed during cif file saving.')
        return ''

    if not report_text:
        document.add_heading('Structure Tables for {}'.format(cif.block.name), 1)
    else:
        document.add_heading('Structure Tables', 1)
        make_columns_section(document, columns='2')

    if report_text:
        if picfile and picfile.exists():
            pic = document.add_paragraph()
            try:
                width = float(options['picture_width'])
            except ValueError:
                width = 7.0
            pic.add_run().add_picture(str(picfile), width=Cm(width))

        paragr = document.add_paragraph()
        paragr.style = document.styles['fliesstext']
        ref = ReferenceList(paragr)
        # -- The main text:
        paragr.add_run('The following text is only a suggestion: ').font.bold = True
        Crystallization(cif, paragr)
        CrstalSelection(cif, paragr)
        MachineType(cif, paragr)
        DataReduct(cif, paragr, ref)
        SpaceChar(paragr).regular()
        SolveRefine(cif, paragr, ref)
        SpaceChar(paragr).regular()
        if cif.hydrogen_atoms_present:
            Hydrogens(cif, paragr)
            SpaceChar(paragr).regular()
        if cif.disorder_present:
            d = Disorder(cif, paragr)
            if d.dsr_sentence:
                ref.append([DSRReference2015(), DSRReference2018()])
                SpaceChar(paragr).regular()
        CCDC(cif, paragr, ref)
        SpaceChar(paragr).regular()
        FinalCifreport(paragr, ref)

    # -- The residuals table:
    table_num = 1
    if report_text:
        # I have to do the header and styling here, otherwise I get another paragraph with a line break in front of the heading.
        p = document.add_paragraph(style='Heading 2')
        p.add_run().add_break(WD_BREAK.COLUMN)
        tab0_head = r"Table {}. Crystal data and structure refinement for {}".format(table_num, cif.block.name)
        p.add_run(text=tab0_head)
    table_num = add_residuals_table(document, cif, table_num)
    p = document.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    if report_text:
        make_columns_section(document, columns='1')
    table_num = add_coords_table(document, cif, table_num)

    if cif.symmops:
        table_num += 1
        document.add_heading(r"Table {}. Bond lengths and angles for {}".format(table_num, cif.block.name), 2)
        make_columns_section(document, columns='2')
        table_num = add_bonds_and_angles_table(document, cif, table_num, without_H)
        make_columns_section(document, columns='1')
        table_num += 1
        document.add_heading(r"Table {}. Torsion angles for {}".format(table_num, cif.block.name), 2)
        make_columns_section(document, columns='2')
        table_num = add_torsion_angles(document, cif, table_num, without_H)
        make_columns_section(document, columns='1')
        if len(list(cif.hydrogen_bonds())) > 0:
            table_num += 1
            document.add_heading(r"Table {}. Hydrogen bonds for {}".format(table_num, cif.block.name), 2)
            table_num = add_hydrogen_bonds(document, cif, table_num)
        document.add_paragraph('')
    else:
        make_columns_section(document, columns='1')
        document.add_paragraph('No further tables, because symmetry operators '
                               '(_space_group_symop_operation_xyz) are missing.')
    if report_text:
        # -- Bibliography:
        document.add_heading('Bibliography', 2)
        ref.make_literature_list(document)

    document.save(output_filename)
    print('\nTables finished - output file: {}'.format(output_filename))
    return file_obj.name


def create_document(report_docx_path: str) -> Document:
    """
    Creates the report docx document.
    :param report_docx_path: Path to the report file.
    :return: The document instance.
    """
    try:
        document = Document(Path(report_docx_path).joinpath(application_path, 'template/template1.docx').absolute())
    except FileNotFoundError as e:
        print(e)
        document = Document()
    # Deleting first (empty) paragraph, otherwise first line would be an empty one:
    try:
        p = document.paragraphs[0]
        delete_paragraph(p)
    except IndexError:
        # no paragraph there
        pass
    return document


def make_columns_section(document, columns: str = '1'):
    """
    Makes a new section (new page) which has a certain number of columns.
    available sections:
    CONTINUOUS, NEW_COLUMN, NEW_PAGE, EVEN_PAGE, ODD_PAGE
    """
    # noinspection PyUnresolvedReferences
    from docx.enum.section import WD_SECTION
    section = document.add_section(WD_SECTION.CONTINUOUS)
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '{}'.format(columns))


def delete_paragraph(paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def set_cell_border(cell: _Cell, **kwargs) -> None:
    """
    Set cell`s border of a Table cell instance.

    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
        bottom={"sz": 12, "color": "#00FF00", "val": "single"},
        start={"sz": 24, "val": "dashed", "shadow": "true"},
        end={"sz": 12, "val": "dashed"},
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # check for tag existnace, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    # list over all available tags
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            # check for tag existnace, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            # looks like order of attributes is important
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def add_residuals_table(document: Document(), cif: CifContainer, table_num: int) -> Tuple[CifContainer, int]:
    # table_num += 1
    exti = cif['_refine_ls_extinction_coef']
    rows = 33
    if cif.is_centrosymm:
        rows -= 1
    # Remove one row for the extinction coefficient:
    if exti in ['.', "'.'", '?', '']:
        rows -= 1
    main_table = document.add_table(rows=rows, cols=2)
    # setup table format:
    set_column_width(main_table.columns[0], Cm(4.05))
    set_column_width(main_table.columns[1], Cm(4.05))
    # Add descriptions to the first column of the main table:
    populate_description_columns(main_table, cif)
    # The main residuals table:
    populate_main_table_values(main_table, cif)
    return table_num


def populate_main_table_values(main_table: Table, cif: CifContainer):
    """
    Fills the main table with residuals. Column, by column.
    """
    header_cells = main_table.rows[0].cells
    header_cells[0].paragraphs[0].add_run('CCDC number')  # .bold = True
    header_cells[1].paragraphs[0].add_run(cif['_database_code_depnum_ccdc_archive'])  # .bold = True

    # Set text for all usual cif keywords by a lookup table:
    for _, key in enumerate(cif_keywords_list):
        # key[1] contains the row number:
        cell = main_table.cell(key[1] + 1, 1)
        if cif[key[0]]:
            cell.text = cif[key[0]]
        else:
            cell.text = '?'
            continue
    # Now the special handling:
    # The sum formula:
    sum_formula = 'no sum formula'
    if cif['_chemical_formula_sum']:
        sum_formula = cif['_chemical_formula_sum']
        ltext2 = sum_formula.replace(" ", "").replace("'", "")
        ltext3 = [''.join(x[1]) for x in it.groupby(ltext2, lambda x: x.isalpha())]
        for _, word in enumerate(ltext3):
            formrun = main_table.cell(1, 1).paragraphs[0]
            formrunsub = formrun.add_run(word)
            if isfloat(word):
                formrunsub.font.subscript = True

    format_space_group(main_table, cif)
    radiation_type = cif['_diffrn_radiation_type']
    radiation_wavelength = cif['_diffrn_radiation_wavelength']
    crystal_size_min = cif['_exptl_crystal_size_min']
    crystal_size_mid = cif['_exptl_crystal_size_mid']
    crystal_size_max = cif['_exptl_crystal_size_max']
    limit_h_min = cif['_diffrn_reflns_limit_h_min']
    limit_h_max = cif['_diffrn_reflns_limit_h_max']
    limit_k_min = cif['_diffrn_reflns_limit_k_min']
    limit_k_max = cif['_diffrn_reflns_limit_k_max']
    theta_min = cif['_diffrn_reflns_theta_min']
    theta_max = cif['_diffrn_reflns_theta_max']
    limit_l_min = cif['_diffrn_reflns_limit_l_min']
    limit_l_max = cif['_diffrn_reflns_limit_l_max']
    reflns_number_total = cif['_reflns_number_total']
    reflns_av_R_equivalents = cif['_diffrn_reflns_av_R_equivalents']
    reflns_av_unetI = cif['_diffrn_reflns_av_unetI/netI']
    ls_number_reflns = cif['_refine_ls_number_reflns']
    ls_number_restraints = cif['_refine_ls_number_restraints']
    ls_number_parameters = cif['_refine_ls_number_parameters']
    ls_R_factor_gt = cif['_refine_ls_R_factor_gt']
    ls_wR_factor_gt = cif['_refine_ls_wR_factor_gt']
    ls_R_factor_all = cif['_refine_ls_R_factor_all']
    ls_wR_factor_ref = cif['_refine_ls_wR_factor_ref']
    goof = cif['_refine_ls_goodness_of_fit_ref']
    try:
        completeness = "{0:.1f} %".format(round(float(cif['_diffrn_measured_fraction_theta_full']) * 100, 1))
    except ValueError:
        completeness = '?'
    try:
        diff_density_min = "{0:.2f}".format(round(float(cif['_refine_diff_density_min']), 2))
    except ValueError:
        diff_density_min = '?'
    try:
        diff_density_max = "{0:.2f}".format(round(float(cif['_refine_diff_density_max']), 2))
    except ValueError:
        diff_density_max = '?'

    # now prepare & write all the concatenated & derived cell contents:
    main_table.cell(17, 1).text = this_or_quest(crystal_size_max) + timessym + \
                                  this_or_quest(crystal_size_mid) + timessym + \
                                  this_or_quest(crystal_size_min)
    wavelength = str(' ({} ='.format(lambdasym) + this_or_quest(radiation_wavelength) +
                     '{}{})'.format(prot_space, angstrom)).replace(' ', '')
    # radtype: ('Mo', 'K', '\\a')
    radtype = format_radiation(radiation_type)
    radrun = main_table.cell(20, 1).paragraphs[0]
    # radiation type e.g. Mo:
    radrun.add_run(radtype[0])
    # K line:
    radrunita = radrun.add_run(radtype[1])
    radrunita.font.italic = True
    alpha = radrun.add_run(radtype[2])
    alpha.font.italic = True
    alpha.font.subscript = True
    # wavelength lambda:
    radrun.add_run(' ' + wavelength)
    try:
        d_max = ' ({:.2f}{}{})'.format(float(radiation_wavelength) / (2 * sin(radians(float(theta_max)))), prot_space,
                                       angstrom)
        # 2theta range:
        main_table.cell(21, 1).text = "{:.2f} to {:.2f}{}".format(2 * float(theta_min), 2 * float(theta_max), d_max)
    except ValueError:
        main_table.cell(21, 1).text = '? to ?'
    main_table.cell(22, 1).text = limit_h_min + ' {} h {} '.format(lessequal, lessequal) + limit_h_max + '\n' \
                                  + limit_k_min + ' {} k {} '.format(lessequal, lessequal) + limit_k_max + '\n' \
                                  + limit_l_min + ' {} l {} '.format(lessequal, lessequal) + limit_l_max
    rint_p = main_table.cell(24, 1).paragraphs[0]
    rint_p.add_run(this_or_quest(reflns_number_total) + '\n')
    rint_p.add_run('R').font.italic = True
    rint_p.add_run('int').font.subscript = True
    rint_p.add_run(' = ' + this_or_quest(reflns_av_R_equivalents) + '\n')
    rint_p.add_run('R').font.italic = True
    rint_p.add_run('sigma').font.subscript = True
    rint_p.add_run(' = ' + this_or_quest(reflns_av_unetI))
    main_table.cell(25, 1).paragraphs[0].add_run(completeness)
    main_table.cell(26, 1).text = this_or_quest(ls_number_reflns) + '/' \
                                  + this_or_quest(ls_number_restraints) + '/' \
                                  + this_or_quest(ls_number_parameters)
    main_table.cell(27, 1).paragraphs[0].add_run(goof)
    r2sig_p = main_table.cell(28, 1).paragraphs[0]
    r2sig_p.add_run('R').font.italic = True
    r2sig_p.add_run('1').font.subscript = True
    r2sig_p.add_run(' = ' + this_or_quest(ls_R_factor_gt))
    r2sig_p.add_run('\nw')
    r2sig_p.add_run('R').font.italic = True
    r2sig_p.add_run('2').font.subscript = True
    r2sig_p.add_run(' = ' + this_or_quest(ls_wR_factor_gt))
    rfull_p = main_table.cell(29, 1).paragraphs[0]
    rfull_p.add_run('R').font.italic = True
    rfull_p.add_run('1').font.subscript = True
    rfull_p.add_run(' = ' + this_or_quest(ls_R_factor_all))
    rfull_p.add_run('\nw')
    rfull_p.add_run('R').font.italic = True
    rfull_p.add_run('2').font.subscript = True
    rfull_p.add_run(' = ' + ls_wR_factor_ref)
    main_table.cell(30, 1).text = diff_density_max + '/' + diff_density_min
    if not cif.is_centrosymm:
        main_table.cell(31, 1).text = cif['_refine_ls_abs_structure_Flack'] or '?'
    exti = cif['_refine_ls_extinction_coef']
    if exti not in ['.', "'.'", '?', '']:
        num = len(main_table.columns[0].cells)
        main_table.columns[1].cells[num - 1].text = exti


def add_decimal_tab(num_string: str) -> str:
    """
    Adds a tab character in front of the decimal point in order to get proper alignment of the tabstops.
    >>> add_decimal_tab('-0.123')
    '-0\\t.123'
    """
    return '\t.'.join(num_string.split('.'))


def add_coords_table(document: Document, cif: CifContainer, table_num: int):
    """
    Adds the table with the atom coordinates.
    :param document: The current word document.
    :param cif: the cif object from CifContainer.
    :return: None
    """
    atoms = list(cif.atoms())
    table_num += 1
    headline = "Table {}. Atomic coordinates and ".format(table_num)
    h = document.add_heading(headline, 2)
    h.add_run('U').font.italic = True
    h.add_run('eq').font.subscript = True
    # eq.italic = True
    h.add_run('{}[{}'.format(prot_space, angstrom))
    h.add_run('2').font.superscript = True
    h.add_run('] for {}'.format(cif.block.name))
    coords_table = document.add_table(rows=len(atoms) + 1, cols=5)
    # coords_table.style = document.styles['Table Grid']
    coords_table.style = 'Table Grid'
    # Atom	x	y	z	U(eq)
    head_row = coords_table.rows[0]
    ar = head_row.cells[0].paragraphs[0].add_run('Atom').bold = True
    px = head_row.cells[1].paragraphs[0]
    ar = px.add_run('x')
    ar.bold = True
    ar.italic = True
    py = head_row.cells[2].paragraphs[0]
    ar = py.add_run('y')
    ar.bold = True
    ar.italic = True
    pz = head_row.cells[3].paragraphs[0]
    ar = pz.add_run('z')
    ar.bold = True
    ar.italic = True
    pu = head_row.cells[4].paragraphs[0]
    ar = pu.add_run('U')
    ar.bold = True
    ar.italic = True
    ar2 = pu.add_run('eq')
    ar2.bold = True
    ar2.font.subscript = True
    # pu.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    # having a list of column cella before is *much* faster!
    col0_cells = coords_table.columns[0].cells
    col1_cells = coords_table.columns[1].cells
    col2_cells = coords_table.columns[2].cells
    col3_cells = coords_table.columns[3].cells
    col4_cells = coords_table.columns[4].cells
    rowidx = 1
    for at in atoms:
        c0, c1, c2, c3, c4 = col0_cells[rowidx], col1_cells[rowidx], col2_cells[rowidx], \
                             col3_cells[rowidx], col4_cells[rowidx]
        rowidx += 1
        c0.text = at[0]  # label
        c1.text = (str(at[2]))  # x
        para_x = c1.paragraphs[0]
        c2.text = (str(at[3]))  # y
        para_y = c2.paragraphs[0]
        c3.text = (str(at[4]))  # z
        para_z = c3.paragraphs[0]
        c4.text = (str(at[7]))  # ueq
        para_u = c4.paragraphs[0]
    p = document.add_paragraph()
    p.style = document.styles['tabunterschr']
    p.add_run('U').font.italic = True
    p.add_run('eq').font.subscript = True
    p.add_run(' is defined as 1/3 of the trace of the orthogonalized ')
    p.add_run('U').font.italic = True
    ij = p.add_run('ij')
    ij.font.subscript = True
    ij.font.italic = True
    p.add_run(' tensor.')
    set_column_width(coords_table.columns[0], Cm(2.3))
    set_column_width(coords_table.columns[1], Cm(2.8))
    set_column_width(coords_table.columns[2], Cm(2.8))
    set_column_width(coords_table.columns[3], Cm(2.8))
    set_column_width(coords_table.columns[4], Cm(2.8))
    document.add_paragraph()
    return table_num


def add_bonds_and_angles_table(document: Document, cif: CifContainer, table_num: int, without_H: bool = False) -> int:
    """
    Make table with bonds and angles.
    """
    # creating rows in advance is *much* faster!
    bond_angle_table = document.add_table(rows=cif.nbonds(without_H) + cif.nangles(without_H) + 3, cols=2,
                                          style='Table Grid')
    # Bond/Angle  value
    head_row = bond_angle_table.rows[0]
    ar = head_row.cells[0].paragraphs[0].add_run('Atom{}Atom'.format(halbgeviert)).bold = True
    p_length = head_row.cells[1].paragraphs[0]
    p_length.add_run('Length [{}]'.format(angstrom)).bold = True
    p_length.paragraph_format.tab_stops.add_tab_stop(Cm(0.4), WD_TAB_ALIGNMENT.DECIMAL, WD_TAB_LEADER.SPACES)
    symms = {}
    newsymms = {}
    num = 1
    # having a list of column cella before is *much* faster!
    col1_cells = bond_angle_table.columns[0].cells
    col2_cells = bond_angle_table.columns[1].cells
    rowidx = 1
    # filling rows:
    for at1, at2, val, symm2 in cif.bonds(without_H):
        c0, c1 = col1_cells[rowidx], col2_cells[rowidx]
        rowidx += 1
        if symm2 == '.':
            symm2 = None
        if symm2 and symm2 not in symms.keys():
            symms[symm2] = num
            # Applys translational symmetry to symmcards:
            # 3_556 -> 2
            card = get_card(cif, symm2)
            s = SymmetryElement(card)
            s.translate(symm2)
            newsymms[num] = s.toShelxl()
            num += 1
        # Atom1 - Atom2:
        c0.text = '{}{}{}'.format(at1, halbgeviert, at2)
        c0.paragraphs[0].add_run('#' + str(symms[symm2]) if symm2 else '').font.superscript = True
        c1.text = str(val)  # bond
    ############ the angles ####################
    head_row = bond_angle_table.rows[rowidx + 1]
    head_row.cells[0].paragraphs[0].add_run('Atom{0}Atom{0}Atom'.format(halbgeviert)).bold = True
    head_row.cells[1].paragraphs[0].add_run('Angle [{}]'.format(degree_sign)).bold = True
    set_cell_border(head_row.cells[0], bottom={"sz": 2, "color": "#000000", "val": "single"})
    set_cell_border(head_row.cells[1], bottom={"sz": 2, "color": "#000000", "val": "single"})
    card = ''
    rowidx += 2
    for at1, at2, at3, angle, symm1, symm3 in cif.angles(without_H):
        c0, c1 = col1_cells[rowidx], col2_cells[rowidx]
        rowidx += 1
        if symm1 == '.':
            symm1 = None
        if symm3 == '.':
            symm3 = None
        if symm1 and symm1 not in symms.keys():
            symms[symm1] = num
            card = get_card(cif, symm1)
            # Applys translational symmetry to symmcards:
            # 3_556 -> 2
            s = SymmetryElement(card)
            s.translate(symm1)
            newsymms[num] = s.toShelxl()
            num += 1
        if symm3 and symm3 not in symms.keys():
            symms[symm3] = num
            card = get_card(cif, symm3)
            s = SymmetryElement(card)
            s.translate(symm3)
            newsymms[num] = s.toShelxl()
            num += 1
        c0.text = at1
        cp0 = c0.paragraphs[0]
        cp0.add_run('#' + str(symms[symm1]) if symm1 else '').font.superscript = True
        cp0.add_run('{}{}{}{}'.format(halbgeviert, at2, halbgeviert, at3))
        cp0.add_run('#' + str(symms[symm3]) if symm3 else '').font.superscript = True
        c1.text = str(angle)  # angle
    set_column_width(bond_angle_table.columns[0], Cm(3.7))
    set_column_width(bond_angle_table.columns[1], Cm(2.5))
    if without_H:
        add_hydrogen_omit_info(document)
    add_last_symminfo_line(newsymms, document)
    return table_num


def add_torsion_angles(document: Document, cif: CifContainer, table_num: int, without_H: bool = False):
    """
    Table 6.  Torsion angles [°] for I-43d_final.
    """
    if not len(list(cif.torsion_angles())) > 0:
        print('No torsion angles in cif.')
        return table_num
    torsion_table = document.add_table(rows=cif.ntorsion_angles(without_H) + 1, cols=2)
    torsion_table.style = 'Table Grid'
    head_row = torsion_table.rows[0]
    head_row.cells[0].paragraphs[0].add_run('Atom{0}Atom{0}Atom{0}Atom'.format(halbgeviert)).bold = True
    head_row.cells[1].paragraphs[0].add_run('Torsion Angle [{}]'.format(degree_sign)).bold = True
    symms = {}
    newsymms = {}
    card = ''
    s = SymmetryElement(card)
    num = 1
    col0_cells = torsion_table.columns[0].cells
    col1_cells = torsion_table.columns[1].cells
    rowidx = 1
    for at1, at2, at3, at4, angle, symm1, symm2, symm3, symm4 in cif.torsion_angles():
        c0, c1 = col0_cells[rowidx], col1_cells[rowidx]
        rowidx += 1
        if symm1 == '.':
            symm1 = None
        if symm2 == '.':
            symm2 = None
        if symm3 == '.':
            symm3 = None
        if symm4 == '.':
            symm4 = None
        if symm1 and symm1 not in symms.keys():
            symms[symm1] = num
            s = SymmetryElement(get_card(cif, symm1))
            s.translate(symm1)
            newsymms[num] = s.toShelxl()
            num += 1
        if symm2 and symm2 not in symms.keys():
            symms[symm2] = num
            s = SymmetryElement(get_card(cif, symm2))
            s.translate(symm2)
            newsymms[num] = s.toShelxl()
            num += 1
        if symm3 and symm3 not in symms.keys():
            symms[symm3] = num
            s = SymmetryElement(get_card(cif, symm3))
            s.translate(symm3)
            newsymms[num] = s.toShelxl()
            num += 1
        if symm4 and symm4 not in symms.keys():
            symms[symm4] = num
            s = SymmetryElement(get_card(cif, symm4))
            s.translate(symm4)
            newsymms[num] = s.toShelxl()
            num += 1
        cp0 = c0.paragraphs[0]
        cp0.add_run(at1)
        cp0.add_run('#' + str(symms[symm1]) if symm1 else '').font.superscript = True
        cp0.add_run(halbgeviert)
        cp0.add_run(at2)
        cp0.add_run('#' + str(symms[symm2]) if symm2 else '').font.superscript = True
        cp0.add_run(halbgeviert)
        cp0.add_run(at3)
        cp0.add_run('#' + str(symms[symm3]) if symm3 else '').font.superscript = True
        cp0.add_run(halbgeviert)
        cp0.add_run(at4)  # labels
        cp0.add_run('#' + str(symms[symm4]) if symm4 else '').font.superscript = True
        c1.paragraphs[0].add_run(str(angle))  # torsion angle
    set_column_width(torsion_table.columns[0], Cm(4.2))
    set_column_width(torsion_table.columns[1], Cm(3.0))
    if without_H:
        add_hydrogen_omit_info(document)
    add_last_symminfo_line(newsymms, document)
    return table_num


def add_hydrogen_bonds(document: Document, cif: CifContainer, table_num: int) -> int:
    """
    Table 7.  Hydrogen bonds for I-43d_final  [Å and °].
    """
    nhydrogenb = len(list(cif.hydrogen_bonds()))
    hydrogen_table = document.add_table(rows=nhydrogenb + 1, cols=5)
    hydrogen_table.style = 'Table Grid'
    head_row = hydrogen_table.rows[0].cells
    # D-H...A	d(D-H)	d(H...A)	d(D...A)	<(DHA)
    head_row[0].paragraphs[0].add_run(
        'D{}H{}A{}[{}]'.format(halbgeviert, ellipsis_mid, prot_space, angstrom)).font.bold = True
    head_row[1].paragraphs[0].add_run('d(D{}H){}[{}]'.format(halbgeviert, prot_space, angstrom)).font.bold = True
    head_row[2].paragraphs[0].add_run('d(H{}A){}[{}]'.format(ellipsis_mid, prot_space, angstrom)).font.bold = True
    head_row[3].paragraphs[0].add_run('d(D{}A){}[{}]'.format(ellipsis_mid, prot_space, angstrom)).font.bold = True
    head_row[4].paragraphs[0].add_run('<(DHA){}[{}]'.format(prot_space, degree_sign)).font.bold = True
    symms = {}
    newsymms = {}
    num = 1
    col0_cells = hydrogen_table.columns[0].cells
    col1_cells = hydrogen_table.columns[1].cells
    col2_cells = hydrogen_table.columns[2].cells
    col3_cells = hydrogen_table.columns[3].cells
    col4_cells = hydrogen_table.columns[4].cells
    rowidx = 1
    for label_d, label_h, label_a, dist_dh, dist_ha, dist_da, angle_dha, symm in cif.hydrogen_bonds():
        c0, c1, c2, c3, c4 = col0_cells[rowidx], col1_cells[rowidx], col2_cells[rowidx], \
                             col3_cells[rowidx], col4_cells[rowidx]
        rowidx += 1
        if symm == '.':
            symm = None
        if symm and symm not in symms.keys():
            symms[symm] = num
            s = SymmetryElement(get_card(cif, symm))
            s.translate(symm)
            newsymms[num] = s.toShelxl()
        num += 1
        symmval = ('#' + str(symms[symm])) if symm else ''
        c0.text = label_d + halbgeviert + label_h + ellipsis_mid + label_a
        c0.paragraphs[0].add_run(symmval).font.superscript = True
        c1.text = dist_dh
        c2.text = dist_ha
        c3.text = dist_da
        c4.text = angle_dha
    widths = (Cm(4), Cm(2.5), Cm(2.5), Cm(2.5), Cm(2.5))
    make_table_widths(hydrogen_table, widths)
    add_last_symminfo_line(newsymms, document)
    return table_num


def populate_description_columns(main_table: Table, cif: CifContainer) -> None:
    """
    This Method adds the descriptions to the fist property table column.
    """
    lgnd1 = main_table.cell(1, 0).paragraphs[0].add_run('Empirical formula')
    lgnd2 = main_table.cell(2, 0).paragraphs[0].add_run('Formula weight')
    lgnd3 = main_table.cell(3, 0).paragraphs[0].add_run('Temperature [K]')
    lgnd4 = main_table.cell(4, 0).paragraphs[0].add_run('Crystal system')
    lgnd5 = main_table.cell(5, 0).paragraphs[0].add_run('Space group (number)')
    lgnd6 = main_table.cell(6, 0).paragraphs[0]
    lgnd6.add_run('a').font.italic = True
    lgnd6.add_run(' [\u00C5]')
    lgnd7 = main_table.cell(7, 0).paragraphs[0]
    lgnd7.add_run('b').font.italic = True
    lgnd7.add_run(' [\u00C5]')
    lgnd8 = main_table.cell(8, 0).paragraphs[0]
    lgnd8.add_run('c').font.italic = True
    lgnd8.add_run(' [\u00C5]')
    lgnd9 = main_table.cell(9, 0).paragraphs[0].add_run('\u03B1 [\u00C5]')
    lgnd10 = main_table.cell(10, 0).paragraphs[0].add_run('\u03B2 [\u00C5]')
    lgnd11 = main_table.cell(11, 0).paragraphs[0].add_run('\u03B3 [\u00C5]')
    lgnd12 = main_table.cell(12, 0).paragraphs[0]
    lgnd12.add_run('Volume [\u00C5')
    lgnd12.add_run('3').font.superscript = True
    lgnd12.add_run(']')
    lgnd13 = main_table.cell(13, 0).paragraphs[0].add_run('Z').font.italic = True
    lgnd14 = main_table.cell(14, 0).paragraphs[0]
    lgnd14.add_run('\u03C1').font.italic = True
    lgnd14.add_run('calc').font.subscript = True
    lgnd14.add_run(' [g/cm')
    lgnd14.add_run('3').font.superscript = True
    lgnd14.add_run(']')
    lgnd15 = main_table.cell(15, 0).paragraphs[0]
    lgnd15.add_run('\u03BC').font.italic = True
    lgnd15.add_run(' [mm')
    lgnd15.add_run('-1').font.superscript = True
    lgnd15.add_run(']')
    lgnd16 = main_table.cell(16, 0).paragraphs[0]
    lgnd16.add_run('F').font.italic = True
    lgnd16.add_run('(000)')
    lgnd17 = main_table.cell(17, 0).paragraphs[0]
    lgnd17.add_run('Crystal size [mm')
    lgnd17.add_run('3').font.superscript = True
    lgnd17.add_run(']')
    lgnd18 = main_table.cell(18, 0).paragraphs[0].add_run('Crystal colour')
    lgnd19 = main_table.cell(19, 0).paragraphs[0].add_run('Crystal shape')
    lgnd20 = main_table.cell(20, 0).paragraphs[0].add_run('Radiation')
    lgnd21 = main_table.cell(21, 0).paragraphs[0].add_run('2\u03F4 range [\u00b0]')
    lgnd22 = main_table.cell(22, 0).paragraphs[0].add_run('Index ranges')
    lgnd23 = main_table.cell(23, 0).paragraphs[0].add_run('Reflections collected')
    lgnd24 = main_table.cell(24, 0).paragraphs[0].add_run('Independent reflections')
    lgnd25 = main_table.cell(25, 0).paragraphs[0]
    theta_full = cif['_diffrn_reflns_theta_full']
    if theta_full:
        lgnd25.add_run('Completeness to \n\u03B8 = {}°'.format(theta_full))
    else:
        lgnd25.add_run('Completeness')
    main_table.cell(26, 0).paragraphs[0].add_run('Data / Restraints / Parameters')
    lgnd27 = main_table.cell(27, 0).paragraphs[0]
    lgnd27.add_run('Goodness-of-fit on ')
    lgnd27.add_run('F').font.italic = True
    lgnd27.add_run('2').font.superscript = True
    lgnd28 = main_table.cell(28, 0).paragraphs[0]
    lgnd28.add_run('Final ')
    lgnd28.add_run('R').font.italic = True
    lgnd28.add_run(' indexes \n[')
    lgnd28.add_run('I').font.italic = True
    lgnd28.add_run('{}2{}('.format(bequal, sigma_sm))
    lgnd28.add_run('I').font.italic = True
    lgnd28.add_run(')]')
    lgnd29 = main_table.cell(29, 0).paragraphs[0]
    lgnd29.add_run('Final ')
    lgnd29.add_run('R').font.italic = True
    lgnd29.add_run(' indexes \n[all data]')
    lgnd30 = main_table.cell(30, 0).paragraphs[0]
    lgnd30.add_run('Largest peak/hole [e{}'.format(angstrom))
    lgnd30.add_run('3').font.superscript = True
    lgnd30.add_run(']')
    if not cif.is_centrosymm:
        lgnd31 = main_table.cell(31, 0).paragraphs[0]
        lgnd31.add_run('Flack X parameter')
    exti = cif['_refine_ls_extinction_coef']
    if exti not in ['.', "'.'", '?', '']:
        # always the last cell
        num = len(main_table.columns[0].cells)
        main_table.columns[0].cells[num - 1].paragraphs[0].add_run('Extinction coefficient')


def set_column_width(column, width: Length) -> None:
    for cell in column.cells:
        cell.width = width


def make_table_widths(table: Table, widths: Sequence[Length]) -> None:
    """
    Sets the width of the columns of a table.
    """
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def add_hydrogen_omit_info(document: Document) -> None:
    """
    Adds information if hydrogen atom bond are omitted from the respective table.
    """
    p = document.add_paragraph('')
    line = 'Bonds to hydrogen atoms were omitted.'
    run = p.add_run(line)
    run.font.size = Pt(8)


def add_last_symminfo_line(newsymms: dict, document: Document) -> None:
    """
    Adds text about the symmetry generators used in order to add symmetry generated atoms.
    """
    p = document.add_paragraph('')
    line = 'Symmetry transformations used to generate equivalent atoms:\n'
    nitems = len(newsymms)
    n = 0
    for key, value in newsymms.items():
        sep = ';'
        if n == nitems:
            sep = ''
        n += 1
        line += "#{}: {}{}   ".format(key, value, sep)
    if newsymms:
        run = p.add_run(line)
        run.font.size = Pt(8)


def get_card(cif: CifContainer, symm: str) -> List[str]:
    """
    Returns a symmetry card from the _space_group_symop_operation_xyz or _symmetry_equiv_pos_as_xyz list.
    :param cif: the cif file object
    :param symm: the symmetry number
    :return: ['x', ' y', ' z'] etc
    """
    card = cif.symmops[int(symm.split('_')[0]) - 1].split(',')
    return card


if __name__ == '__main__':
    output_filename = 'tables.docx'
    import time

    # make_report_from(get_files_from_current_dir()[5])
    t1 = time.perf_counter()
    report_options = {'report_text': True,
                      'picture_width': 7.5,
                      'without_H': False,
                      }
    make_report_from(options=report_options, file_obj=Path(r'tests/examples/1979688.cif'),
                     output_filename=Path(output_filename),
                     picfile=Path('icon/finalcif.png'))
    # make_report_from(Path(r'/Volumes/nifty/p-1.cif'))
    t2 = time.perf_counter()
    print('complete table:', round(t2 - t1, 2), 's')
    if os.name == 'nt':
        subprocess.call(['cmd', '/C', Path(output_filename).absolute()])
    else:
        subprocess.call(['open', Path(output_filename).absolute()])
    # make_report_from(Path(r'test-data/sad-final.cif'))
    # make_report_from(Path(r'/Volumes/home/strukturen/eigene/DK_30011/sad-final.cif'))
    # make_report_from(Path(r'D:\goedaten\strukturen_goe\eigene\DK_4008\xl12\new\r3c.cif'))
