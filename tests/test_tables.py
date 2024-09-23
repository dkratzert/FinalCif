import os
import unittest
from pathlib import Path
from unittest.mock import Mock

import docx
from docx import Document
from docx.enum.shape import WD_INLINE_SHAPE
from docx.shape import InlineShapes
from docx.shared import Cm
from docx.table import Table
from packaging.version import Version

from finalcif.appwindow import AppWindow
from finalcif.cif.cif_file_io import CifContainer
from finalcif.report.tables import make_report_from

os.environ["RUNNING_TEST"] = 'True'
data = Path('tests')
test_data = Path('test-data')


class TablesTestCase(unittest.TestCase):

    def setUp(self) -> None:
        self.testcif = (data / 'examples/1979688_small.cif').absolute()
        self.myapp = AppWindow(file=self.testcif)
        self.myapp.ui.HAtomsCheckBox.setChecked(False)
        self.myapp.ui.ReportTextCheckBox.setChecked(False)
        self.myapp.ui.PictureWidthDoubleSpinBox.setValue(0.0)
        # make sure to use no template:
        self.myapp.ui.docxTemplatesListWidget.setCurrentRow(0)
        # self.myapp.show()
        self.reportdoc = self.myapp.cif.finalcif_file_prefixed(prefix='report_', suffix='-finalcif.docx')
        self.report_zip = self.myapp.cif.finalcif_file_prefixed(prefix='', suffix='-finalcif.zip')
        self.myapp.ui.PictureWidthDoubleSpinBox.setValue(7.43)
        self.myapp.select_report_picture(Path('finalcif/icon/finalcif.png'))

    def tearDown(self) -> None:
        self.myapp.cif.finalcif_file.unlink(missing_ok=True)
        self.reportdoc.unlink(missing_ok=True)
        self.report_zip.unlink(missing_ok=True)
        self.myapp.ui.ReportTextCheckBox.setChecked(False)
        self.myapp.ui.HAtomsCheckBox.setChecked(False)
        self.myapp.ui.PictureWidthDoubleSpinBox.setValue(7.5)
        self.myapp.close()

    def test_picture_has_correct_size(self):
        self.myapp.ui.SaveFullReportButton.click()
        doc = Document(self.reportdoc.absolute())
        shapes: InlineShapes = doc.inline_shapes
        self.assertEqual(WD_INLINE_SHAPE.PICTURE, shapes[0].type)
        self.assertEqual(Cm(7.43).emu, shapes[0].width)

    def test_default_picture_width(self):
        self.myapp.ui.PictureWidthDoubleSpinBox.setValue(0.0)
        self.myapp.ui.SaveFullReportButton.click()
        doc = Document(self.reportdoc.resolve())
        shapes: InlineShapes = doc.inline_shapes
        self.assertEqual(WD_INLINE_SHAPE.PICTURE, shapes[0].type)
        self.assertEqual(Cm(7.5).emu, shapes[0].width)


class TemplateReportWithoutAppTestCase(unittest.TestCase):
    def setUp(self) -> None:
        self.testcif = (data / 'examples/1979688.cif').absolute()
        self.cif = CifContainer(self.testcif)
        self.options = Mock()
        self.options.picture_width = 7.43
        self.options.without_h = False
        self.text_template = Path('finalcif/template/template_text.docx').absolute()
        self.template_without_text = Path('finalcif/template/template_without_text.docx').absolute()
        self.reportdoc = self.cif.finalcif_file_prefixed(prefix='report_', suffix='-finalcif.docx')
        self.report_zip = self.cif.finalcif_file_prefixed(prefix='', suffix='-finalcif.zip')
        self.report_pic = Path('finalcif/icon/finalcif.png')

    def tearDown(self) -> None:
        self.reportdoc.unlink(missing_ok=True)
        self.report_zip.unlink(missing_ok=True)

    def test_option_with_h(self):
        make_report_from(options=self.options, cif=self.cif,
                         output_filename=str(self.reportdoc), picfile=self.report_pic)
        doc = Document(self.reportdoc.absolute())
        table: Table = doc.tables[3]
        self.assertEqual('C1–H1', table.cell(row_idx=4, col_idx=0).text)

    def test_option_without_h(self):
        self.options.without_h = True
        make_report_from(options=self.options, cif=self.cif,
                         output_filename=str(self.reportdoc), picfile=self.report_pic)
        doc = Document(self.reportdoc.absolute())
        table: Table = doc.tables[3]
        self.assertEqual('O1–C13', table.cell(row_idx=4, col_idx=0).text)

    def test_all_paragraphs(self):
        make_report_from(options=self.options, cif=self.cif,
                         output_filename=str(self.reportdoc), picfile=self.report_pic)
        doc = Document(self.reportdoc.absolute())
        newline = '\n' if Version(docx.__version__) < Version('1.0') else ''
        result = ('Structure Tables\n'
                  '\n'
                  '\n'
                  '\n'
                  f'{newline}Table 1. Crystal data and structure refinement for cu_BruecknerJK_153F40_0m\n{newline}'
                  '\n'
                  '\n'
                  'Refinement details for cu_BruecknerJK_153F40_0m\n'
                  'The methanol molecule is disordered around a special position and thus half '
                  'occupied.\n'
                  'Table 2. Atomic coordinates and Ueq\xa0[Å2] for cu_BruecknerJK_153F40_0m\n'
                  'Ueq is defined as 1/3 of the trace of the orthogonalized Uij tensor.\n'
                  '\n'
                  'Table 3. Anisotropic displacement parameters [Å2] for '
                  'cu_BruecknerJK_153F40_0m.\n'
                  'The anisotropic displacement factor exponent takes the form: −2π2[\u205f'
                  'h2(a*)2U11\u205f+\u205fk2(b*)2U22\u205f+\u205f…\u205f+\u205f2hka*b*U12\u205f'
                  ']\n'
                  'Table 4. Bond lengths and angles for cu_BruecknerJK_153F40_0m\n'
                  '\n'
                  '\n'
                  'Table 5. Torsion angles for cu_BruecknerJK_153F40_0m\n'
                  '\n'
                  '\n'
                  '\n'
                  'Bibliography'
                  )

        self.assertEqual(result, '\n'.join([x.text for x in doc.paragraphs]))


class TablesNoPictureTestCase(unittest.TestCase):
    def setUp(self) -> None:
        self.testcif = (data / 'examples/1979688.cif').absolute()
        self.cif = CifContainer(self.testcif)
        self.options = Mock()
        self.options.picture_width = 7.43
        self.options.without_h = False
        self.text_template = Path('finalcif/template/template_text.docx').absolute()
        self.template_without_text = Path('finalcif/template/template_without_text.docx').absolute()
        self.reportdoc = self.cif.finalcif_file_prefixed(prefix='report_', suffix='-finalcif.docx')
        self.report_zip = self.cif.finalcif_file_prefixed(prefix='', suffix='-finalcif.zip')
        self.report_pic = Path('finalcif/icon/finalcif.png')

    def tearDown(self) -> None:
        self.reportdoc.unlink(missing_ok=True)
        self.report_zip.unlink(missing_ok=True)

    def test_save_report_works(self):
        make_report_from(options=self.options, cif=self.cif,
                         output_filename=str(self.reportdoc), picfile=None)
        self.assertEqual(True, self.reportdoc.exists())

    def test_picture_has_correct_size(self):
        make_report_from(options=self.options, cif=self.cif,
                         output_filename=str(self.reportdoc), picfile=None)
        doc = Document(self.reportdoc.resolve())
        shapes: InlineShapes = doc.inline_shapes
        self.assertEqual(0, len(shapes))

    def test_picture_shape_exists(self):
        make_report_from(options=self.options, cif=self.cif,
                         output_filename=str(self.reportdoc), picfile=self.report_pic)
        doc = Document(self.reportdoc.resolve())
        shapes: InlineShapes = doc.inline_shapes
        self.assertEqual(1, len(shapes))


class ReportWithsymmetryTestCase(unittest.TestCase):
    def setUp(self) -> None:
        self.testcif = (test_data / 'p31c.cif').resolve()
        self.cif = CifContainer(self.testcif)
        self.options = Mock()
        self.options.without_h = False
        self.reportdoc = self.cif.finalcif_file_prefixed(prefix='report_', suffix='-finalcif.docx')

    def tearDown(self) -> None:
        self.reportdoc.unlink(missing_ok=True)
        self.cif.finalcif_file.unlink(missing_ok=True)

    def test_symmetry_indicators(self):
        make_report_from(options=self.options, cif=self.cif, output_filename=str(self.reportdoc), picfile=None)
        doc = Document(self.reportdoc.absolute())
        table: Table = doc.tables[3]
        # Bond:
        self.assertEqual('C2–C3#1', table.cell(row_idx=18, col_idx=0).text)
        # Angle:
        self.assertEqual("C3'#1–C2'–C3'", table.cell(row_idx=148, col_idx=0).text)
        # Torsion angle:
        table: Table = doc.tables[4]
        self.assertEqual("C3#1–C2–C3–N1", table.cell(row_idx=6, col_idx=0).text)
        # Hydrogen bond:
        table: Table = doc.tables[5]
        self.assertEqual("N1^a–H1^a⋯Cl1#1", table.cell(row_idx=1, col_idx=0).text)


if __name__ == '__main__':
    unittest.main()
