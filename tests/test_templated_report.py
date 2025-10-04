import os
from unittest.mock import Mock

from finalcif.tools.options import Options

os.environ['RUNNING_TEST'] = 'True'
import unittest
from pathlib import Path

from docx import Document
from docx.enum.shape import WD_INLINE_SHAPE
from docx.shape import InlineShapes
from docx.shared import Cm
from docx.table import Table

from finalcif.appwindow import AppWindow
from finalcif.cif.cif_file_io import CifContainer
from finalcif.report.templated_report import TemplatedReport, ReportFormat, Hydrogens

data = Path('tests')
test_data = Path('test-data')


# noinspection PyMissingTypeHints
class TemplateReportTestCase(unittest.TestCase):
    def setUp(self) -> None:
        self.testcif = (data / 'examples/1979688.cif').absolute()
        self.myapp = AppWindow(file=self.testcif)
        self.myapp.ui.HAtomsCheckBox.setChecked(False)
        self.myapp.ui.ReportTextCheckBox.setChecked(False)
        self.myapp.ui.PictureWidthDoubleSpinBox.setValue(7.43)
        self.options = Mock()
        self.options.picture_width = 7.43
        self.options.without_h = False
        self.options.use_picometers = False
        self.text_template = Path('finalcif/template/report_default.docx').absolute()
        self.template_without_text = Path('finalcif/template/template_without_text.docx').absolute()
        self.import_templates()
        self.myapp.ui.docxTemplatesListWidget.setCurrentRow(2)
        self.reportdoc = self.myapp.cif.finalcif_file_prefixed(prefix='report_', suffix='-finalcif.docx')
        self.report_zip = self.myapp.cif.finalcif_file_prefixed(prefix='', suffix='-finalcif.zip')
        self.report_pic = Path('finalcif/icon/finalcif.png')
        self.myapp.select_report_picture(self.report_pic)

    def tearDown(self) -> None:
        self.myapp.cif.finalcif_file.unlink(missing_ok=True)
        self.reportdoc.unlink(missing_ok=True)
        self.report_zip.unlink(missing_ok=True)
        self.myapp.ui.ReportTextCheckBox.setChecked(False)
        self.myapp.ui.HAtomsCheckBox.setChecked(False)
        self.myapp.ui.PictureWidthDoubleSpinBox.setValue(7.5)
        self.myapp.ui.docxTemplatesListWidget.blockSignals(True)
        self._clean_templates()
        self.myapp.ui.docxTemplatesListWidget.blockSignals(False)
        self.myapp.close()

    def import_templates(self):
        # blocking signals, because signal gets fired after delete and crashes: 
        self.myapp.ui.docxTemplatesListWidget.blockSignals(True)
        self._clean_templates()
        self.myapp.templates.add_new_template(str(self.text_template))
        self.myapp.templates.add_new_template(str(self.template_without_text))
        print('imported templates')
        self.myapp.ui.docxTemplatesListWidget.blockSignals(False)

    def _clean_templates(self):
        for num in range(1, self.myapp.ui.docxTemplatesListWidget.count()):
            self.myapp.ui.docxTemplatesListWidget.setCurrentRow(num)
            self.myapp.templates.remove_current_template()

    def test_with_report_text(self):
        self.import_templates()
        self.myapp.ui.docxTemplatesListWidget.setCurrentRow(2)
        self.myapp.ui.SaveFullReportButton.click()
        doc = Document(self.reportdoc.absolute())
        # for n, p in enumerate(doc.paragraphs):
        #    print(n, p.text)
        self.assertEqual('The compound was crystalli', doc.paragraphs[2].text[:26])


class TemplateReportWithoutAppTestCase(unittest.TestCase):
    def setUp(self) -> None:
        self.testcif = (data / 'examples/1979688.cif').absolute()
        cif = CifContainer(self.testcif)
        self.options = Options()
        self.options._picture_width = 7.43
        self.options._without_h = False
        self.text_template = Path('finalcif/template/report_default.docx').absolute()
        self.template_without_text = Path('finalcif/template/template_without_text.docx').absolute()
        self.reportdoc = cif.finalcif_file_prefixed(prefix='report_', suffix='-finalcif.docx')
        self.report_zip = cif.finalcif_file_prefixed(prefix='', suffix='-finalcif.zip')
        self.report_pic = Path('finalcif/icon/finalcif.png')

    def tearDown(self) -> None:
        self.reportdoc.unlink(missing_ok=True)
        self.report_zip.unlink(missing_ok=True)

    def test_ccdc_num_in_table(self):
        t = TemplatedReport(options=self.options,
                            cif=CifContainer(self.testcif),
                            format=ReportFormat.RICHTEXT)
        ok = t.make_templated_docx_report(output_filename=str(self.reportdoc),
                                          picfile=self.report_pic,
                                          template_path=self.text_template)
        self.assertTrue(ok)
        doc = Document(str(self.reportdoc.absolute()))
        table: Table = doc.tables[0]
        # This is with the 'CCDC' string, because CCDC will be deleted during CIF save in  main application:
        self.assertEqual('1979688', table.cell(row_idx=0, col_idx=1).text)

    def test_picture_has_correct_size(self):
        """
        For this test, self.myapp.set_report_picture(Path('finalcif/icon/finalcif.png'))
        has to be set correctly.
        """
        t = TemplatedReport(options=self.options,
                            cif=CifContainer(self.testcif),
                            format=ReportFormat.RICHTEXT)
        ok = t.make_templated_docx_report(output_filename=str(self.reportdoc),
                                          picfile=self.report_pic,
                                          template_path=self.text_template)
        self.assertTrue(ok)
        doc = Document(self.reportdoc.absolute())
        shapes: InlineShapes = doc.inline_shapes
        self.assertEqual(WD_INLINE_SHAPE.PICTURE, shapes[0].type)
        self.assertEqual(Cm(7.43).emu, shapes[0].width)

    def test_citations(self):
        t = TemplatedReport(options=self.options,
                            cif=CifContainer(self.testcif),
                            format=ReportFormat.RICHTEXT)
        t.make_templated_docx_report(output_filename=str(self.reportdoc),
                                     picfile=self.report_pic,
                                     template_path=self.text_template)
        doc = Document(self.reportdoc.absolute())
        paragraphs = [p.text for p in doc.paragraphs]
        self.assertTrue('Bibliography' in paragraphs)


class TestReportFromMultiCif(unittest.TestCase):
    def setUp(self):
        self.reportdoc = Path('test.docx')
        self.reportdoc.unlink(missing_ok=True)
        self.docx_templ = test_data / 'templates/test_template_for_multitable_dist.docx'
        self.multi_cif = test_data / '1000007-multi.cif'
        self.cif = CifContainer(file=self.multi_cif)
        self.options = Options()

    def tearDown(self):
        self.reportdoc.unlink(missing_ok=True)

    def test_get_distance_from_atoms(self):
        self.options._without_h = False
        self.options._picture_width = 7.43
        t = TemplatedReport(options=self.options,
                            cif=self.cif,
                            format=ReportFormat.RICHTEXT)
        ok = t.make_templated_docx_report(output_filename='test.docx', picfile=Path(),
                                          template_path=self.docx_templ)
        self.assertTrue(ok)
        doc = Document(self.reportdoc.resolve().__str__())
        self.assertEqual('C1-C2 in p-1 distance: 1.5123(17)', doc.paragraphs[0].text)
        self.assertEqual('C1-C2 in p21c distance: 1.544(3)', doc.paragraphs[1].text)


class TestCIFwithOneAtom(unittest.TestCase):

    def setUp(self) -> None:
        # creating a new CIF with a new block:
        self.cif = CifContainer('test-data/diamond/9008564.cif')
        self.options = Options()
        self.t = TemplatedReport(options=self.options,
                                 cif=self.cif,
                                 format=ReportFormat.HTML)

    def test_foo(self):
        context = self.t.get_context(cif=self.cif, options=self.options, picfile=Path())
        self.assertEqual(({'label': 'C',
                           'occ'  : '1.000000',
                           'part' : '0',
                           'type' : 'C',
                           'u_eq' : '.',
                           'x'    : '0.00000',
                           'y'    : '0.00000',
                           'z'    : '0.00000'},), context.get('atomic_coordinates'))
        self.assertEqual([], context.get('bonds'))
        self.assertEqual([], context.get('angles'))
        self.assertEqual((), context.get('displacement_parameters'))
        self.assertEqual('??', context.get('solution_method'))
        self.assertEqual('??', context.get('solution_program'))
        self.assertEqual('', context.get('resolution_angstrom'))
        self.assertEqual('Å', context.get('dist_unit'))
        self.assertEqual('', context.get('completeness'))
        self.assertEqual([], context.get('torsions'))
        self.assertEqual('?', context.get('t_min'))
        self.assertEqual('', context.get('ba_symminfo'))
        self.assertEqual(('<i>F</i><i>d</i><span style=" text-decoration: overline;">3</span><i>m</i> '
                          '(227)'), context.get('space_group'))


class TestData(unittest.TestCase):

    def setUp(self) -> None:
        # creating a new CIF with a new block:
        self.cif = CifContainer('foo.cif', 'testblock')
        self.options = Options()
        self.t = TemplatedReport(options=self.options,
                                 cif=self.cif,
                                 format=ReportFormat.RICHTEXT)

    def test_get_integration_program_with_spaces(self):
        # Here we have the special case, that there is no space character between the version number and the bracket.
        self.cif['_computing_data_reduction'] = 'CrysAlisPro 1.171.39.20a (Rigaku OD, 2015)'
        self.assertEqual('CrysAlisPro 1.171.39.20a (Rigaku OD, 2015)', self.cif['_computing_data_reduction'])
        result = self.t.text_formatter.get_integration_program(self.cif)
        self.assertEqual('CrysAlisPro', result)
        self.assertEqual('Crysalispro, 1.171.39.20a, 2015, Rigaku OD.',
                         str(self.t.text_formatter.literature['integration']))

    def test_get_integration_program_with_line_break(self):
        # Here we have the special case, that there is no space character between the version number and the bracket.
        self.cif['_computing_data_reduction'] = "CrysAlisPro 1.171.39.20a\n" \
                                                "(Rigaku OD, 2015)\n"
        self.assertEqual('CrysAlisPro 1.171.39.20a\n(Rigaku OD, 2015)\n', self.cif['_computing_data_reduction'])
        result = self.t.text_formatter.get_integration_program(self.cif)
        self.assertEqual('CrysAlisPro', result)
        self.assertEqual('Crysalispro, 1.171.39.20a, 2015, Rigaku OD.',
                         str(self.t.text_formatter.literature['integration']))

    def test_get_integration_program_with_missing_information(self):
        # Here we have all in one line with spaces inbetween:
        self.cif['_computing_data_reduction'] = 'CrysAlisPro 1.171.39.20a (Rigaku OD)'
        self.assertEqual('CrysAlisPro 1.171.39.20a (Rigaku OD)', self.cif['_computing_data_reduction'])
        result = self.t.text_formatter.get_integration_program(self.cif)
        self.assertEqual('CrysAlisPro', result)
        self.assertEqual('Crysalispro, unknown version, Rigaku OD.',
                         str(self.t.text_formatter.literature['integration']))

    def test_get_integration_program_saint(self):
        # Here we have all in one line with spaces inbetween:
        self.cif['_computing_data_reduction'] = 'SAINT V8.40A'
        result = self.t.text_formatter.get_integration_program(self.cif)
        self.assertEqual('SAINT V8.40A', result)
        self.assertEqual('Bruker, SAINT, V8.40A, Bruker AXS Inc., Madison, Wisconsin, USA.',
                         str(self.t.text_formatter.literature['integration']))

    def test_get_integration_program_saint_without_version(self):
        # Here we have all in one line with spaces inbetween:
        self.cif['_computing_data_reduction'] = 'SAINT'
        result = self.t.text_formatter.get_integration_program(self.cif)
        self.assertEqual('SAINT', result)
        self.assertEqual('Bruker, SAINT, Bruker AXS Inc., Madison, Wisconsin, USA.',
                         str(self.t.text_formatter.literature['integration']))


class TestHydrogenText(unittest.TestCase):
    def setUp(self) -> None:
        self.cif = CifContainer('test-data/p21c.cif')
        self.h = Hydrogens(self.cif)

    def test_hydrogen_text_html(self):
        result = ('All hydrogen atoms were refined isotropic on calculated positions using a '
                  'riding model with their <i>U</i><sub>iso</sub> values constrained to 1.5 '
                  'times the <i>U</i><sub>eq</sub> of their pivot atoms for terminal '
                  'sp<sup>3</sup> carbon atoms and 1.2 times for all other carbon atoms.')
        self.assertEqual(result, self.h.html())

    def test_hydrogen_text_richtext(self):
        result = ('<w:r><w:t xml:space="preserve">All hydrogen atoms were refined isotropic '
                  '</w:t></w:r><w:r><w:t xml:space="preserve">on calculated positions using a '
                  'riding model with their </w:t></w:r><w:r><w:rPr><w:i/></w:rPr><w:t '
                  'xml:space="preserve">U</w:t></w:r><w:r><w:rPr><w:vertAlign '
                  'w:val="subscript"/></w:rPr><w:t '
                  'xml:space="preserve">iso</w:t></w:r><w:r><w:t xml:space="preserve"> values '
                  'constrained to 1.5 times the </w:t></w:r><w:r><w:rPr><w:i/></w:rPr><w:t '
                  'xml:space="preserve">U</w:t></w:r><w:r><w:rPr><w:vertAlign '
                  'w:val="subscript"/></w:rPr><w:t xml:space="preserve">eq</w:t></w:r><w:r><w:t '
                  'xml:space="preserve"> of their pivot atoms for terminal '
                  'sp</w:t></w:r><w:r><w:rPr><w:vertAlign w:val="superscript"/></w:rPr><w:t '
                  'xml:space="preserve">3</w:t></w:r><w:r><w:t xml:space="preserve"> carbon '
                  'atoms and 1.2 times for all other carbon atoms</w:t></w:r><w:r><w:t xml:space="preserve">'
                  '. </w:t></w:r>')
        self.assertEqual(result, self.h.richtext().xml)


if __name__ == '__main__':
    unittest.main()
